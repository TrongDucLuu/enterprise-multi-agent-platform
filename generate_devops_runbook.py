#!/usr/bin/env python3
"""
Script to generate DEVOPS_RUNBOOK.docx
Comprehensive DevOps Runbook & Operational Guide for IT Helpdesk Multi-Agent AI System.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_styler import (
    NAVY_PRIMARY,
    BLUE_SECONDARY,
    TEXT_DARK,
    GRAY_MUTED,
    WHITE,
    COLOR_CODE_BG,
    COLOR_CALLOUT_NOTE_BG,
    COLOR_CALLOUT_NOTE_BORDER,
    COLOR_CALLOUT_WARN_BG,
    COLOR_CALLOUT_WARN_BORDER,
    COLOR_CALLOUT_IMPORTANT_BG,
    COLOR_CALLOUT_IMPORTANT_BORDER,
    COLOR_CALLOUT_TIP_BG,
    COLOR_CALLOUT_TIP_BORDER,
    COLOR_TABLE_HEADER_BG,
    COLOR_TABLE_ROW_ALT,
    COLOR_TABLE_BORDER,
    set_cell_background,
    set_cell_margins,
    set_table_borders,
    add_styled_heading,
    add_body_paragraph,
    add_code_block,
    add_callout,
    format_custom_table,
)


def build_runbook_docx(output_path):
    """Main builder function for the complete DevOps Runbook."""
    doc = Document()
    
    # Set page margins to 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header setup
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("IT Helpdesk Multi-Agent AI | Enterprise DevOps Runbook v2.1")
        hr.font.name = "Arial"
        hr.font.size = Pt(8.5)
        hr.font.color.rgb = GRAY_MUTED
        
        # Footer setup
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("CONFIDENTIAL - INTERNAL IT & DEVOPS USE ONLY")
        fr.font.name = "Arial"
        fr.font.size = Pt(8.5)
        fr.font.color.rgb = GRAY_MUTED

    # --- DOCUMENT HEADER / TITLE ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("HƯỚNG DẪN VẬN HÀNH & TRIỂN KHAI HẠ TẦNG\n(DEVOPS RUNBOOK & OPERATIONAL GUIDE)")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = NAVY_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Hệ thống IT Helpdesk Multi-Agent AI Enterprise (Google Cloud Run, BigQuery Vector Search, ADK & SSO)")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = BLUE_SECONDARY

    # Metadata Table
    meta_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        meta_table,
        [1.5, 1.75, 1.5, 1.75],
        ["Thuộc tính", "Giá trị", "Thuộc tính", "Giá trị"],
        [
            ["Hệ thống", "IT Helpdesk Multi-Agent AI", "Phiên bản tài liệu", "2.1.0-Enterprise"],
            ["Môi trường mục tiêu", "Production / Staging / Dev", "Chủ quản kỹ thuật", "DevOps & Cloud Platform"],
            ["Mô hình AI", "Gemini 3 Flash & Pro", "Phê duyệt bởi", "Principal Cloud Architect"],
            ["Hạ tầng Cloud", "Google Cloud Platform (GCP)", "Ngày cập nhật", "28/08/2026"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- MỤC 1: TỔNG QUAN KIẾN TRÚC HẠ TẦNG ---
    add_styled_heading(doc, "1. TỔNG QUAN KIẾN TRÚC HẠ TẦNG & THÀNH PHẦN VẬN HÀNH", 1)
    
    add_body_paragraph(
        doc,
        "Hệ thống IT Helpdesk Multi-Agent AI được thiết kế theo kiến trúc phi trạng thái (Stateless Microservices), "
        "đóng gói container chuẩn OCI và triển khai trên hạ tầng Google Cloud Run thế hệ 2 (Cloud Run v2). "
        "Hệ thống phối hợp chặt chẽ giữa các dịch vụ PaaS và Serverless cao cấp của Google Cloud Platform nhằm đảm bảo "
        "tối ưu hóa độ trễ, khả năng tự động mở rộng (Autoscaling) và tối thiểu hóa chi phí vận hành cố định."
    )

    add_body_paragraph(
        doc,
        "",
        bold_prefix="Danh mục các thành phần hạ tầng cốt lõi:"
    )

    infra_table = doc.add_table(rows=1, cols=3)
    format_custom_table(
        infra_table,
        [1.8, 1.8, 2.9],
        ["Thành phần GCP", "Dịch vụ đảm nhiệm", "Mô tả vai trò kỹ thuật & Cơ chế vận hành"],
        [
            ["Compute Layer", "Google Cloud Run (v2)", "Chạy ứng dụng FastAPI & ADK Agents. Cấu hình 2 vCPU, 2GB RAM, concurrency 80, tự động scale 0 -> N instances."],
            ["Vector Knowledge Base", "BigQuery Vector Search", "Lưu trữ tài liệu RAG & vector embeddings (Dataset: it_helpdesk_kb). Serverless 100%, chi phí cố định 0 USD."],
            ["Ticket State Store", "Cloud Firestore", "Lưu trữ trạng thái Helpdesk Tickets (Collection: helpdesk_tickets). Write-through caching kết hợp in-memory fallback."],
            ["Identity & SSO", "Google OIDC + Cloud IAM", "Xác thực danh tính qua JWKS public keys của accounts.google.com, phân quyền theo email domain và vai trò RBAC."],
            ["Semantic Cache", "In-Memory Cosine Cache", "Bộ đệm ngữ nghĩa câu hỏi lặp lại với ngưỡng Cosine Sim >= 0.92, phản hồi < 50ms, tiết kiệm 100% token Gemini."],
            ["Secret Management", "GCP Secret Manager", "Lưu trữ an toàn các biến môi trường nhạy cảm, API keys và secret JWT mà không hardcode vào mã nguồn."],
            ["Container Registry", "Artifact Registry", "Kho lưu trữ Docker images bảo mật (Format: Docker v2, Region: us-central1 hoặc asia-southeast1)."],
            ["Observability", "Cloud Trace & Logging", "Thu thập structured logs chuẩn JSON, tracing phân tán qua OpenTelemetry và giám sát sức khỏe dịch vụ."],
        ]
    )

    add_callout(
        doc,
        "TIP",
        "Tối ưu hóa Chi phí Hạ tầng Serverless",
        "Nhờ áp dụng BigQuery Vector Search thay thế cho Vertex AI Vector Search Index Endpoints chuyên dụng, chi phí hạ tầng tĩnh giảm từ ~300 USD/tháng xuống gần 0 USD/tháng cho các doanh nghiệp có kho tri thức dưới 100.000 vectors."
    )

    # --- MỤC 2: BẢNG QUẢN LÝ BIẾN MÔI TRƯỜNG & BÍ MẬT BẢO MẬT ---
    add_styled_heading(doc, "2. BẢNG QUẢN LÝ BIẾN MÔI TRƯỜNG & BÍ MẬT BẢO MẬT", 1)
    
    add_body_paragraph(
        doc,
        "Mọi tham số cấu hình của hệ thống được quản lý thông qua biến môi trường (Environment Variables) "
        "và Google Secret Manager. Khi triển khai trên Production, toàn bộ các biến bảo mật phải được nạp thông qua cơ chế "
        "Secret Reference của Cloud Run thay vì gán giá trị Plaintext."
    )

    env_table = doc.add_table(rows=1, cols=5)
    format_custom_table(
        env_table,
        [1.6, 0.9, 1.2, 0.8, 2.0],
        ["Tên Biến Môi Trường", "Bắt buộc", "Giá trị mặc định", "Secret Mngr", "Mục đích & Hướng dẫn Cấu hình"],
        [
            ["ENVIRONMENT", "Có", "development", "Không", "Môi trường thực thi: 'development', 'staging', 'production'."],
            ["GOOGLE_CLOUD_PROJECT", "Có", "—", "Không", "ID của dự án Google Cloud Platform (GCP Project ID)."],
            ["GOOGLE_CLOUD_REGION", "Có", "us-central1", "Không", "Vùng triển khai Cloud Run và BigQuery (vd: us-central1)."],
            ["SSO_CLIENT_ID", "Có", "—", "Có", "OAuth 2.0 Client ID được cấp từ GCP Identity Console."],
            ["ALLOWED_DOMAINS", "Có (Prod)", "—", "Không", "Danh sách domain email công ty được phép (vd: company.com,corp.com). Bắt buộc phải set trong Prod để kích hoạt Fail-Closed."],
            ["ALLOW_LOCAL_DEV_SSO", "Không", "false", "Không", "Bật tính năng giả lập SSO cho local dev. Tự động bị khóa thành false trên môi trường Production."],
            ["SSO_JWT_SECRET", "Không", "—", "Có", "Khóa bí mật HS256 chỉ dùng để ký token thử nghiệm trong local dev."],
            ["USE_FIRESTORE_TICKETS", "Không", "false", "Không", "Bật lưu trữ Firestore cho tickets (tự động bật trên Cloud Run)."],
            ["KNOWLEDGE_BACKEND", "Không", "in_memory", "Không", "Backend cho RAG Knowledge Store: 'in_memory' hoặc 'bigquery'."],
            ["BIGQUERY_KB_DATASET", "Không", "it_helpdesk_kb", "Không", "Dataset BigQuery chứa bảng vectors và tài liệu tri thức."],
            ["SEMANTIC_CACHE_ENABLED", "Không", "true", "Không", "Bật/Tắt lớp bộ đệm ngữ nghĩa cho câu hỏi lặp lại."],
            ["SEMANTIC_CACHE_THRESHOLD", "Không", "0.92", "Không", "Ngưỡng độ tương đồng Cosine Similarity để coi là trùng khớp."],
            ["OTEL_TO_CLOUD", "Không", "false", "Không", "Đẩy OpenTelemetry traces lên GCP Cloud Trace."],
        ]
    )

    add_callout(
        doc,
        "WARNING",
        "Quy định An toàn Thông tin - Biến ALLOWED_DOMAINS",
        "Trên môi trường Production, nếu ALLOWED_DOMAINS bị bỏ trống, hệ thống sẽ kích hoạt cơ chế Fail-Closed và từ chối toàn bộ token đăng nhập. Tuyệt đối không để rỗng cấu hình này khi bàn giao cho bộ phận Vận hành."
    )

    # --- MỤC 3: HƯỚNG DẪN KHỞI TẠO HẠ TẦNG TỰ ĐỘNG (TERRAFORM) ---
    add_styled_heading(doc, "3. HƯỚNG DẪN KHỞI TẠO HẠ TẦNG TỰ ĐỘNG (TERRAFORM IAC)", 1)
    
    add_body_paragraph(
        doc,
        "Toàn bộ tài nguyên đám mây được định nghĩa dưới dạng mã (Infrastructure as Code) trong thư mục deployment/terraform. "
        "Bộ mã Terraform đã được cấu hình tuân thủ nguyên tắc đặc quyền tối thiểu (Least Privilege) và quản lý vòng đời tài nguyên thông minh."
    )

    add_styled_heading(doc, "3.1. Cấu trúc Tài nguyên Terraform (Terraform Resource Mapping)", 2)
    
    add_body_paragraph(
        doc,
        "• google_service_account.helpdesk_sa: Service Account đại diện cho ứng dụng chạy trên Cloud Run.\n"
        "• google_project_iam_member: Cấp các quyền chặt chẽ: roles/aiplatform.user (Vertex AI), roles/bigquery.dataEditor & roles/bigquery.jobUser (BigQuery), roles/logging.logWriter (Cloud Logging), roles/secretmanager.secretAccessor (Secret Manager).\n"
        "• google_bigquery_dataset.kb_dataset: Khởi tạo dataset it_helpdesk_kb lưu trữ vector articles.\n"
        "• google_artifact_registry_repository.app_repo: Kho chứa Docker Image định dạng Docker v2.\n"
        "• google_cloud_run_v2_service.helpdesk_service: Định nghĩa Cloud Run Service với cấu hình CPU, Memory và Environment Variables."
    )

    add_styled_heading(doc, "3.2. Quy trình Thực thi Terraform", 2)
    
    add_body_paragraph(
        doc,
        "Thực hiện theo các bước sau từ máy trạm DevOps hoặc qua Terraform Cloud / CI Runner:"
    )

    add_code_block(
        doc,
        "# Bước 1: Di chuyển vào thư mục Terraform\n"
        "cd deployment/terraform\n\n"
        "# Bước 2: Khởi tạo Terraform Provider & State Backend\n"
        "terraform init\n\n"
        "# Bước 3: Kiểm tra kế hoạch thay đổi (Dry-Run)\n"
        "terraform plan \\\n"
        "  -var=\"project_id=my-company-it-prod\" \\\n"
        "  -var=\"region=us-central1\" \\\n"
        "  -var=\"sso_client_id=123456789-abc.apps.googleusercontent.com\" \\\n"
        "  -var=\"allowed_domains=mycompany.com,corp.mycompany.com\" \\\n"
        "  -out=tfplan.binary\n\n"
        "# Bước 4: Áp dụng thay đổi vào hạ tầng GCP\n"
        "terraform apply tfplan.binary"
    )

    add_callout(
        doc,
        "IMPORTANT",
        "Cơ chế Chống Drift Cấu hình CI/CD (Lifecycle Ignore Changes)",
        "Trong file main.tf, khối cấu hình Cloud Run đã được gán 'lifecycle { ignore_changes = [template[0].containers[0].image] }'. Điều này đảm bảo khi Pipeline CI/CD cập nhật SHA của image container mới, Terraform sẽ không tự ý rollback image về phiên bản cũ khi chạy lệnh apply định kỳ."
    )

    # --- MỤC 4: QUY TRÌNH ĐÓNG GÓI CONTAINER & TRIỂN KHAI ---
    add_styled_heading(doc, "4. QUY TRÌNH ĐÓNG GÓI CONTAINER & TRIỂN KHAI (CI/CD PIPELINE)", 1)
    
    add_body_paragraph(
        doc,
        "Hệ thống sử dụng Dockerfile chuẩn Production với chiến lược Multi-Stage Build và quản lý package cực nhanh qua công cụ UV của Astral. "
        "Container được chạy dưới quyền người dùng không đặc quyền (non-root user `appuser` với UID 1001) để phòng chống leo thang đặc quyền."
    )

    add_styled_heading(doc, "4.1. Quy trình Build & Push Container Image", 2)

    add_code_block(
        doc,
        "# Đặt biến môi trường dự án\n"
        "export PROJECT_ID=\"my-company-it-prod\"\n"
        "export REGION=\"us-central1\"\n"
        "export REPO=\"it-helpdesk-repo\"\n"
        "export IMAGE_TAG=\"$(git rev-parse --short HEAD)\"\n"
        "export IMAGE_URI=\"${REGION}-docker.pkg.dev/${PROJECT_ID}/${REPO}/it-helpdesk-agent:${IMAGE_TAG}\"\n\n"
        "# Build và push image lên Artifact Registry qua Cloud Build\n"
        "gcloud builds submit --tag ${IMAGE_URI} .\n\n"
        "# Hoặc sử dụng Makefile tích hợp sẵn\n"
        "make docker-deploy PROJECT_ID=${PROJECT_ID} REGION=${REGION}"
    )

    add_styled_heading(doc, "4.2. Cập nhật Revision Mới trên Cloud Run", 2)

    add_code_block(
        doc,
        "# Cập nhật container image cho Cloud Run service\n"
        "gcloud run services update it-helpdesk-agent \\\n"
        "  --image=${IMAGE_URI} \\\n"
        "  --region=${REGION} \\\n"
        "  --project=${PROJECT_ID}"
    )

    add_styled_heading(doc, "4.3. Chiến lược Di chuyển Lưu lượng & Rollback (Traffic Migration)", 2)
    
    add_body_paragraph(
        doc,
        "Cloud Run tự động kích hoạt Revision mới với 100% traffic sau khi vượt qua Healthcheck. "
        "Nếu phát hiện lỗi nghiêm trọng, thực hiện rollback về Revision trước đó chỉ trong vòng dưới 30 giây bằng lệnh:"
    )

    add_code_block(
        doc,
        "# Liệt kê các revision gần nhất\n"
        "gcloud run revisions list --service=it-helpdesk-agent --region=${REGION}\n\n"
        "# Điều hướng 100% lưu lượng về Revision ổn định trước đó (vd: it-helpdesk-agent-00042-xyz)\n"
        "gcloud run services update-traffic it-helpdesk-agent \\\n"
        "  --to-revisions=it-helpdesk-agent-00042-xyz=100 \\\n"
        "  --region=${REGION}"
    )

    # --- MỤC 5: QUY TRÌNH VẬN HÀNH CHUẨN (SOPS) ---
    add_styled_heading(doc, "5. QUY TRÌNH VẬN HÀNH CHUẨN (STANDARD OPERATING PROCEDURES)", 1)

    add_styled_heading(doc, "SOP-01: Đồng bộ Môi trường & Dependencies Cục bộ", 3)
    add_body_paragraph(
        doc,
        "Sử dụng công cụ UV để khởi tạo và đồng bộ virtual environment:"
    )
    add_code_block(doc, "uv sync")

    add_styled_heading(doc, "SOP-02: Khởi chạy Chế độ Kiểm thử Tương tác CLI", 3)
    add_body_paragraph(
        doc,
        "Cho phép kỹ sư DevOps kiểm tra trực tiếp phản hồi của Orchestrator và các Sub-agent từ Terminal:"
    )
    add_code_block(doc, "uv run python main.py --mode cli")

    add_styled_heading(doc, "SOP-03: Khởi chạy Web Server FastAPI & ADK UI", 3)
    add_body_paragraph(
        doc,
        "Khởi chạy web server phục vụ API và giao diện tương tác:"
    )
    add_code_block(doc, "uv run python main.py --mode serve --host 0.0.0.0 --port 8080")

    add_styled_heading(doc, "SOP-04: Thực thi Toàn bộ Bộ Kiểm thử Tự động (Pytest)", 3)
    add_body_paragraph(
        doc,
        "Bắt buộc thực thi trước khi commit code hoặc deploy lên bất kỳ môi trường nào. Bộ kiểm thử bao gồm 46 test cases:"
    )
    add_code_block(doc, "uv run pytest tests/ -v")

    add_styled_heading(doc, "SOP-05: Quản trị & Giám sát Semantic Cache", 3)
    add_body_paragraph(
        doc,
        "Kiểm tra hiệu quả tỷ lệ Cache Hit Rate thông qua API Endpoint:"
    )
    add_code_block(
        doc,
        "# Lấy thống kê hiệu năng bộ đệm\n"
        "curl -X GET http://localhost:8080/api/cache/stats \\\n"
        "  -H \"Authorization: Bearer <VALID_OIDC_TOKEN>\"\n\n"
        "# Tra cứu trực tiếp câu hỏi tương tự\n"
        "curl -X GET \"http://localhost:8080/api/cache/query?q=cách+đổi+mật+khẩu+wifi&threshold=0.92\" \\\n"
        "  -H \"Authorization: Bearer <VALID_OIDC_TOKEN>\""
    )

    add_styled_heading(doc, "SOP-06: Quy trình Onboarding Khách Hàng Mới & Nạp Knowledge Base", 3)
    add_body_paragraph(
        doc,
        "Quy trình 5 bước nạp tài liệu và cấu hình hệ thống nghiệp vụ cho khách hàng doanh nghiệp mới mà không cần sửa code:\n"
        "1. Khai báo hệ thống nghiệp vụ trong `config/systems.yaml` (ví dụ: ERP, HRM, CRM, CORE_BANKING) và các vai trò quản trị tương ứng.\n"
        "2. Đặt tài liệu kỹ thuật/SOP (.md, .txt, .docx, .jsonl) vào thư mục `data/knowledge_base/`.\n"
        "3. Chạy kiểm thử mô phỏng (Dry-Run) để xác thực định dạng và phân tách đoạn (semantic chunking):\n"
        "   `make ingest-kb-dry`\n"
        "4. Thực hiện nạp chính thức dữ liệu và sinh vector embedding 768-dim (text-embedding-005) vào BigQuery:\n"
        "   `make ingest-kb`\n"
        "5. Xác minh truy vấn ngữ nghĩa với cờ `--test-query`:\n"
        "   `python scripts/ingest_knowledge_base.py --test-query \"lỗi phân quyền SAP PO\"`"
    )

    # --- MỤC 6: GIÁM SÁT, CẢNH BÁO & KHẢ NĂNG QUAN SÁT ---
    add_styled_heading(doc, "6. GIÁM SÁT, CẢNH BÁO & KHẢ NĂNG QUAN SÁT (OBSERVABILITY & MONITORING)", 1)
    
    add_body_paragraph(
        doc,
        "Hệ thống cung cấp khả năng quan sát toàn diện thông qua 3 trụ cột: Metrics, Logs và Traces."
    )

    add_styled_heading(doc, "6.1. Healthcheck & Liveness/Readiness Probes", 2)
    add_body_paragraph(
        doc,
        "Endpoint `/healthz` hoạt động công khai không cần auth, phục vụ kiểm tra sức khỏe của Cloud Run Load Balancer:\n"
        "• Trả về `{\"status\": \"ok\", \"service\": \"it_helpdesk_agent\"}` với mã HTTP 200 OK."
    )

    add_styled_heading(doc, "6.2. Ma trận Cảnh báo Vận hành (Alerting Policy Matrix)", 2)

    alert_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        alert_table,
        [1.6, 1.4, 1.5, 2.0],
        ["Chỉ số Giám sát", "Ngưỡng Cảnh báo", "Mức độ Nghiêm trọng", "Hành động Vận hành Khuyến nghị"],
        [
            ["HTTP 5xx Error Rate", "> 1.0% trong 5 phút", "CRITICAL (P1)", "Kiểm tra Cloud Logging tìm Exception. Rollback revision nếu xảy ra sau khi deploy."],
            ["P95 Latency", "> 3.0 giây trong 5 phút", "HIGH (P2)", "Kiểm tra tỉ lệ Cache Hit Rate, độ trễ API Vertex AI hoặc tắc nghẽn BigQuery."],
            ["OIDC Auth Failures", "> 50 lỗi / phút", "HIGH (P2)", "Kiểm tra cấu hình ALLOWED_DOMAINS, JWKS Google certs hoặc token Client ID."],
            ["Container Memory Usage", "> 85% giới hạn (1.7GB)", "MEDIUM (P3)", "Kiểm tra rò rỉ bộ nhớ (Memory Leak), điều chỉnh giới hạn RAM lên 4GB."],
            ["BigQuery Quota Error", "> 5 lỗi QuotaExceeded", "MEDIUM (P3)", "Yêu cầu tăng quota BigQuery Query Execution trên GCP Console."],
        ]
    )

    # --- MỤC 7: KỊCH BẢN XỬ LÝ SỰ CỐ & KHÔI PHỤC THẢM HỌA ---
    add_styled_heading(doc, "7. KỊCH BẢN XỬ LÝ SỰ CỐ & KHÔI PHỤC THẢM HỌA (INCIDENT RESPONSE)", 1)

    add_styled_heading(doc, "Sự cố 1: Lỗi Xác thực SSO Hàng loạt (HTTP 401/403)", 2)
    add_body_paragraph(
        doc,
        "• Triệu chứng: Người dùng không thể đăng nhập hoặc nhận thông báo 'Truy cập bị từ chối: Domain không được phép'.\n"
        "• Nguyên nhân gốc rễ: Biến ALLOWED_DOMAINS bị thiếu domain mới sáp nhập, hoặc SSO_CLIENT_ID không khớp với Google Client ID trên Frontend.\n"
        "• Quy trình xử lý:\n"
        "  1. Kiểm tra log Cloud Logging: `resource.type=\"cloud_run_revision\" textPayload=~\"OIDC\"`.\n"
        "  2. Cập nhật biến môi trường trên Cloud Run:\n"
        "     `gcloud run services update it-helpdesk-agent --set-env-vars=\"ALLOWED_DOMAINS=company.com,newdomain.com\"`"
    )

    add_styled_heading(doc, "Sự cố 2: Lỗi Kết nối Cloud Firestore", 2)
    add_body_paragraph(
        doc,
        "• Triệu chứng: Log xuất hiện 'Firestore unavailable... Falling back to in-memory ticketing store'.\n"
        "• Tác động: Dịch vụ không bị sập (nhờ cơ chế fallback tự động), nhưng dữ liệu ticket mới sẽ lưu tạm trên RAM và mất khi container restart.\n"
        "• Quy trình xử lý:\n"
        "  1. Kiểm tra quyền IAM của Service Account: đảm bảo có quyền `roles/datastore.user`.\n"
        "  2. Kiểm tra Firestore Database đã được kích hoạt trên GCP Project hay chưa."
    )

    add_styled_heading(doc, "Sự cố 3: Vượt Hạn Ngạch Vertex AI (Rate Limit 429)", 2)
    add_body_paragraph(
        doc,
        "• Triệu chứng: Log xuất hiện mã lỗi `RESOURCE_EXHAUSTED` hoặc `429 Too Many Requests`.\n"
        "• Cơ chế tự phục hồi: Mã nguồn đã tích hợp `HttpRetryOptions(attempts=3)` với Exponential Backoff.\n"
        "• Quy trình xử lý khẩn cấp: Tăng hạn ngạch Quota 'GenerateContent requests per minute' trên Google Cloud Console Quotas & Limits."
    )

    # --- MỤC 8: DANH MỤC KIỂM TRA BẢO MẬT & BÀN GIAO VẬN HÀNH ---
    add_styled_heading(doc, "8. DANH MỤC KIỂM TRA BẢO MẬT & BÀN GIAO (GO-LIVE CHECKLIST)", 1)

    checklist_table = doc.add_table(rows=1, cols=3)
    format_custom_table(
        checklist_table,
        [2.5, 1.2, 2.8],
        ["Hạng mục Kiểm tra", "Trạng thái", "Yêu cầu Tiêu chuẩn Bắt buộc"],
        [
            ["HTTPS / SSL Termination", "BẮT BUỘC", "Cloud Run tự động kích hoạt chứng chỉ SSL/TLS được quản lý bởi Google."],
            ["Google OIDC JWKS Verification", "BẮT BUỘC", "Đảm bảo verify chữ ký RS256 trực tiếp từ accounts.google.com."],
            ["Fail-Closed Domain Whitelist", "BẮT BUỘC", "Biến ALLOWED_DOMAINS đã cấu hình chính xác danh sách domain nội bộ."],
            ["Dev Mock Token Disabled", "BẮT BUỘC", "Biến ALLOW_LOCAL_DEV_SSO phải bằng 'false' trong môi trường Production."],
            ["Secrets in Secret Manager", "BẮT BUỘC", "Tuyệt đối không lưu API keys hoặc secrets dưới dạng Plaintext trong mã nguồn."],
            ["Service Account Least Privilege", "BẮT BUỘC", "Chỉ cấp các IAM roles tối thiểu cần thiết (Vertex, BigQuery, Logging, Secrets)."],
            ["Unit Tests 100% Pass", "BẮT BUỘC", "Bộ 46 unit test cases phải vượt qua 100% trước khi bàn giao."],
        ]
    )

    add_body_paragraph(doc, "")
    add_body_paragraph(
        doc,
        "Tài liệu này là quy chuẩn vận hành chính thức và bắt buộc tuân thủ đối với toàn bộ kỹ sư DevOps, "
        "SRE và Cloud Platform tham gia quản trị hệ thống IT Helpdesk Multi-Agent AI.",
        bold_prefix="KẾT LUẬN & CAM KẾT VẬN HÀNH: "
    )

    # Save to path
    doc.save(output_path)
    print(f"Successfully generated DevOps Runbook: {output_path}")

if __name__ == "__main__":
    target_path = "/Users/luuduc/.gemini/antigravity/scratch/it-helpdesk-agent/DEVOPS_RUNBOOK.docx"
    build_runbook_docx(target_path)
