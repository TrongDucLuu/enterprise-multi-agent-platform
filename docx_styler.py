#!/usr/bin/env python3
"""
Shared Styling and Component Utilities for Document Generation (docx)
Used by generate_srs_document.py and generate_devops_runbook.py.
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- Shared Color Palette Constants ---
NAVY_PRIMARY = RGBColor(26, 54, 93)      # #1A365D - Main Headings
BLUE_SECONDARY = RGBColor(43, 108, 176)   # #2B6CB0 - Sub-headings
TEXT_DARK = RGBColor(45, 55, 72)         # #2D3748 - Body Text
GRAY_MUTED = RGBColor(113, 128, 150)     # #718096 - Metadata / Captions
WHITE = RGBColor(255, 255, 255)

COLOR_CODE_BG = "F1F5F9"                 # Light Slate Gray for code
COLOR_CALLOUT_NOTE_BG = "EBF8FF"         # Soft Blue
COLOR_CALLOUT_NOTE_BORDER = "3182CE"
COLOR_CALLOUT_WARN_BG = "FFF5F5"         # Soft Red
COLOR_CALLOUT_WARN_BORDER = "E53E3E"
COLOR_CALLOUT_IMPORTANT_BG = "FFFAF0"    # Soft Orange/Amber
COLOR_CALLOUT_IMPORTANT_BORDER = "DD6B20"
COLOR_CALLOUT_TIP_BG = "F0FFF4"          # Soft Green
COLOR_CALLOUT_TIP_BORDER = "38A169"
COLOR_TABLE_HEADER_BG = "1E3A8A"         # Deep Blue
COLOR_TABLE_ROW_ALT = "F8FAFC"           # Very Light Slate
COLOR_TABLE_BORDER = "CBD5E1"            # Light Border


def set_cell_background(cell, fill_hex: str):
    """Sets cell shading color in hex."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top: int = 140, bottom: int = 140, left: int = 180, right: int = 180):
    """Sets cell padding (in dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_table_borders(table, border_color: str = "CBD5E1"):
    """Sets subtle table borders."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_styled_heading(doc, text: str, level: int):
    """Adds styled headings with proper spacing and color."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run.font.size = Pt(15)
        run.font.color.rgb = NAVY_PRIMARY
        run.font.name = "Arial"
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(12.5)
        run.font.color.rgb = BLUE_SECONDARY
        run.font.name = "Arial"
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY_PRIMARY
        run.font.name = "Arial"
    return p


def add_body_paragraph(doc, text: str = "", space_after: int = 6, bold_prefix: str = None):
    """Adds a standardized body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.size = Pt(10)
        r_prefix.font.color.rgb = TEXT_DARK
        r_prefix.font.name = "Arial"
        
    if text:
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = TEXT_DARK
        r.font.name = "Arial"
    return p


def add_code_block(doc, code_text: str):
    """Creates a distinct monospace code block with shaded background and border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, COLOR_CODE_BG)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Border for the code box
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="94A3B8"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)


def add_callout(doc, callout_type: str, title: str, message: str):
    """Adds a formatted callout box (NOTE, WARNING, IMPORTANT, TIP)."""
    type_configs = {
        "NOTE": (COLOR_CALLOUT_NOTE_BG, COLOR_CALLOUT_NOTE_BORDER, "📌 GHI CHÚ (NOTE):"),
        "WARNING": (COLOR_CALLOUT_WARN_BG, COLOR_CALLOUT_WARN_BORDER, "⚠️ CẢNH BÁO (WARNING):"),
        "IMPORTANT": (COLOR_CALLOUT_IMPORTANT_BG, COLOR_CALLOUT_IMPORTANT_BORDER, "❗ QUAN TRỌNG (IMPORTANT):"),
        "TIP": (COLOR_CALLOUT_TIP_BG, COLOR_CALLOUT_TIP_BORDER, "💡 GỢI Ý KỸ THUẬT (TIP):"),
    }
    bg_color, border_color, default_prefix = type_configs.get(
        callout_type.upper(), 
        (COLOR_CALLOUT_NOTE_BG, COLOR_CALLOUT_NOTE_BORDER, "📌 NOTE:")
    )
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{default_prefix} {title}\n" if title else f"{default_prefix}\n")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = NAVY_PRIMARY
    
    r_msg = p.add_run(message)
    r_msg.font.name = "Arial"
    r_msg.font.size = Pt(9.5)
    r_msg.font.color.rgb = TEXT_DARK
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)


def format_custom_table(table, col_widths: list[float], headers: list[str], data: list[list]):
    """Fills and formats a professional table with custom header and alternate row shading."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Inches(col_widths[i])
        set_cell_background(cell, COLOR_TABLE_HEADER_BG)
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        
    # Set cantSplit and tblHeader
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    # Format Data Rows
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        bg_color = COLOR_TABLE_ROW_ALT if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = Inches(col_widths[col_idx])
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(str(text))
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT_DARK
