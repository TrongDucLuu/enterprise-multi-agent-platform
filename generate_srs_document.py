#!/usr/bin/env python3
"""
Script to generate PRODUCT_SRS_DOCUMENT.docx
Enterprise-grade Product Requirements & Software Requirements Specification (SRS)
for IT Helpdesk Multi-Agent AI System.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- Color Palette Constants ---
NAVY_PRIMARY = RGBColor(26, 54, 93)      # #1A365D - Main Headings
BLUE_SECONDARY = RGBColor(43, 108, 176)   # #2B6CB0 - Sub-headings
TEXT_DARK = RGBColor(45, 55, 72)         # #2D3748 - Body Text
GRAY_MUTED = RGBColor(113, 128, 150)     # #718096 - Metadata / Captions
WHITE = RGBColor(255, 255, 255)
COLOR_CODE_BG = "F1F5F9"                 # Light Slate Gray for code
COLOR_CALLOUT_NOTE_BG = "EBF8FF"         # Soft Blue
COLOR_CALLOUT_NOTE_BORDER = "3182CE"
COLOR_CALLOUT_WARN_BG = "FFF5F5"         # Soft Red
COLOR_CALLOUT_WARN_BORDER = "E53E3E"
COLOR_CALLOUT_IMPORTANT_BG = "FFFAF0"    # Soft Orange/Amber
COLOR_CALLOUT_IMPORTANT_BORDER = "DD6B20"
COLOR_CALLOUT_TIP_BG = "F0FFF4"          # Soft Green
COLOR_CALLOUT_TIP_BORDER = "38A169"
COLOR_TABLE_HEADER_BG = "1E3A8A"         # Deep Navy Blue
COLOR_TABLE_ROW_ALT = "F8FAFC"           # Very Light Slate
COLOR_TABLE_BORDER = "CBD5E1"            # Light Border

def set_cell_background(cell, fill_hex):
    """Sets cell shading color in hex."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Sets cell padding (in dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, border_color="CBD5E1"):
    """Sets subtle table borders."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_heading(doc, text, level):
    """Adds styled headings with proper spacing and color."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run.font.size = Pt(15)
        run.font.color.rgb = NAVY_PRIMARY
        run.font.name = "Arial"
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(12.5)
        run.font.color.rgb = BLUE_SECONDARY
        run.font.name = "Arial"
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY_PRIMARY
        run.font.name = "Arial"
    return p

def add_body_paragraph(doc, text="", space_after=6, bold_prefix=None):
    """Adds a standardized body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.size = Pt(10)
        r_prefix.font.color.rgb = TEXT_DARK
        r_prefix.font.name = "Arial"
        
    if text:
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = TEXT_DARK
        r.font.name = "Arial"
    return p

def add_code_block(doc, code_text):
    """Creates a distinct monospace code block with shaded background and border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, COLOR_CODE_BG)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Border for the code box
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="94A3B8"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_callout(doc, callout_type, title, message):
    """Adds a formatted callout box (NOTE, WARNING, IMPORTANT, TIP)."""
    type_configs = {
        "NOTE": (COLOR_CALLOUT_NOTE_BG, COLOR_CALLOUT_NOTE_BORDER, "📌 GHI CHÚ (NOTE):"),
        "WARNING": (COLOR_CALLOUT_WARN_BG, COLOR_CALLOUT_WARN_BORDER, "⚠️ CẢNH BÁO (WARNING):"),
        "IMPORTANT": (COLOR_CALLOUT_IMPORTANT_BG, COLOR_CALLOUT_IMPORTANT_BORDER, "❗ QUAN TRỌNG (IMPORTANT):"),
        "TIP": (COLOR_CALLOUT_TIP_BG, COLOR_CALLOUT_TIP_BORDER, "💡 GỢI Ý KỸ THUẬT (TIP):"),
    }
    bg_color, border_color, default_prefix = type_configs.get(
        callout_type.upper(), 
        (COLOR_CALLOUT_NOTE_BG, COLOR_CALLOUT_NOTE_BORDER, "📌 NOTE:")
    )
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{default_prefix} {title}\n" if title else f"{default_prefix}\n")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = NAVY_PRIMARY
    
    r_msg = p.add_run(message)
    r_msg.font.name = "Arial"
    r_msg.font.size = Pt(9.5)
    r_msg.font.color.rgb = TEXT_DARK
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def format_custom_table(table, col_widths, headers, data):
    """Fills and formats a professional table with custom header and alternate row shading."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Inches(col_widths[i])
        set_cell_background(cell, COLOR_TABLE_HEADER_BG)
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        
    # Set cantSplit and tblHeader
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    # Format Data Rows
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        bg_color = COLOR_TABLE_ROW_ALT if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = Inches(col_widths[col_idx])
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(str(text))
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT_DARK

def build_srs_docx(output_path):
    """Main builder function for the complete Product & SRS Document."""
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("IT Helpdesk Multi-Agent AI | Enterprise Product Specification & SRS")
        hr.font.name = "Arial"
        hr.font.size = Pt(8.5)
        hr.font.color.rgb = GRAY_MUTED
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("CONFIDENTIAL - ENTERPRISE BUSINESS & TECHNICAL REQUIREMENTS")
        fr.font.name = "Arial"
        fr.font.size = Pt(8.5)
        fr.font.color.rgb = GRAY_MUTED

    # --- COVER / TITLE SECTION ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("TÀI LIỆU ĐẶC TẢ YÊU CẦU SẢN PHẨM & NGHIỆP VỤ\n(PRODUCT REQUIREMENTS & SOFTWARE SPECIFICATION - SRS)")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(17)
    run_title.font.color.rgb = NAVY_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Hệ thống Trợ lý Trí tuệ Nhân tạo Đa Tác tử Hỗ trợ Vận hành Kỹ thuật Doanh nghiệp (IT Helpdesk Multi-Agent AI)")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = BLUE_SECONDARY

    # Metadata Table
    meta_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        meta_table,
        [1.5, 1.75, 1.5, 1.75],
        ["Thuộc tính", "Giá trị", "Thuộc tính", "Giá trị"],
        [
            ["Tên Sản phẩm", "IT Helpdesk Multi-Agent AI", "Phiên bản SRS", "2.1.0-Enterprise"],
            ["Loại Tài liệu", "Product Spec & Functional SRS", "Chủ quản Sản phẩm", "Lead PM & Senior BA"],
            ["Đối tượng Áp dụng", "Toàn bộ Doanh nghiệp / Khối IT", "Kiến trúc Nền tảng", "Google ADK & Gemini 3"],
            ["Trạng thái Phê duyệt", "Approved for Production", "Ngày ban hành", "28/08/2026"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- MỤC 1: TỔNG QUAN SẢN PHẨM & TẦM NHÌN CHIẾN LƯỢC ---
    add_styled_heading(doc, "1. TỔNG QUAN SẢN PHẨM & TẦM NHÌN CHIẾN LƯỢC (PRODUCT OVERVIEW)", 1)
    
    add_body_paragraph(
        doc,
        "1.1. Bối cảnh & Thách thức Doanh nghiệp (Business Problem Statement):",
        bold_prefix=""
    )
    add_body_paragraph(
        doc,
        "Tại các doanh nghiệp hiện đại, khối vận hành IT Helpdesk đóng vai trò sống còn để duy trì năng suất lao động "
        "cho hàng ngàn nhân viên. Tuy nhiên, các mô hình Helpdesk truyền thống đang gặp phải các điểm nghẽn nghiêm trọng:\n"
        "• 60% – 70% yêu cầu gửi đến là các tác vụ lặp lại cơ bản (FAQ, reset mật khẩu, mở khóa tài khoản, cài đặt Wi-Fi/VPN).\n"
        "• Thời gian phản hồi ban đầu (First Response Time - FRT) và thời gian giải quyết sự cố (MTTR) kéo dài từ 4 đến 24 giờ.\n"
        "• Sự cố nghiệp vụ trên các hệ thống ERP, HRM, CRM đòi hỏi tra cứu tài liệu hướng dẫn kỹ thuật dài và phân tán.\n"
        "• Sự cố gián đoạn hệ thống nghiêm trọng (Downtime) thiếu công cụ tự động phân tích nhanh nguyên nhân gốc rễ (Root Cause Analysis - RCA).\n"
        "• Chi phí nhân sự IT Support gia tăng tuyến tính theo quy mô nhân sự của công ty."
    )

    add_body_paragraph(
        doc,
        "1.2. Tầm nhìn Sản phẩm (Product Vision):",
        bold_prefix=""
    )
    add_body_paragraph(
        doc,
        "Xây dựng một nền tảng **Trợ lý IT Helpdesk AI Đa Tác tử thông minh, tự chủ và an toàn cấp Doanh nghiệp**, "
        "đóng vai trò là điểm tiếp nhận duy nhất (Single Point of Contact - SPOC), có khả năng giải quyết tự động tức thì các sự cố "
        "thông thường, hỗ trợ tra cứu tri thức nghiệp vụ sâu, đồng thời đồng hành cùng kỹ sư IT cấp cao trong việc chẩn đoán lỗi hạ tầng "
        "và giám sát tuân thủ cam kết SLA của các nhà cung cấp dịch vụ."
    )

    add_body_paragraph(
        doc,
        "1.3. Mục tiêu Nghiệp vụ & Chỉ số Đánh giá Hiệu quả (Business Goals & OKRs/KPIs):",
        bold_prefix=""
    )

    kpi_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        kpi_table,
        [1.6, 1.4, 1.5, 2.0],
        ["Chỉ số Đo lường (KPI)", "Trước khi áp dụng", "Mục tiêu sau triển khai", "Ý nghĩa Nghiệp vụ & Đóng góp"],
        [
            ["Tỷ lệ Tự phục vụ Thành công (FCR)", "< 15%", ">= 65%", "Người dùng tự xử lý thành công ngay ở Mức 1 mà không cần can thiệp con người."],
            ["Thời gian Giải quyết Sự cố (MTTR)", "4.5 giờ", "< 15 phút (Trung bình)", "Giảm thiểu thời gian gián đoạn công việc của nhân viên văn phòng."],
            ["Tỷ lệ Phản hồi Tức thì (Cache Hit)", "0%", ">= 50% câu hỏi lặp", "Lớp Semantic Cache phản hồi < 50ms, tiết kiệm 100% token Gemini."],
            ["Độ chính xác Báo cáo L3 RCA", "Thủ công (1–2 ngày)", "Tức thì (< 30 giây)", "Phân loại chính xác 6 nhóm lỗi cốt lõi (OOM, DB, Disk, Null, Auth, Network)."],
            ["Chỉ số Hài lòng Người dùng (CSAT)", "3.2 / 5.0", ">= 4.7 / 5.0", "Trải nghiệm hội thoại tự nhiên, thông minh, phục vụ 24/7/365."],
        ]
    )

    # --- MỤC 2: CHÂN DUNG NGƯỜI DÙNG & CÁC BÊN LIÊN QUAN ---
    add_styled_heading(doc, "2. CHÂN DUNG NGƯỜI DÙNG & BÊN LIÊN QUAN (USER PERSONAS & STAKEHOLDERS)", 1)

    persona_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        persona_table,
        [1.5, 1.4, 1.8, 1.8],
        ["Chân dung (Persona)", "Vai trò Doanh nghiệp", "Nhu cầu & Nỗi đau chính", "Kỳ vọng đối với Hệ thống AI"],
        [
            ["Nhân viên Văn phòng\n(Employee)", "Người dùng cuối\n(End-User)", "Quên mật khẩu, lỗi Wi-Fi, không vào được tài khoản, cần hỗ trợ gấp.", "Nhận hướng dẫn chi tiết từng bước, tự mở khóa tài khoản tức thì, tạo ticket tự động."],
            ["Chuyên viên IT Helpdesk\n(IT Support / L1-L2)", "Hỗ trợ Kỹ thuật\n(IT Specialist)", "Quá tải các ticket lặp lại, mất thời gian tra cứu manual ERP/HRM/CRM.", "AI tự động phân loại, tóm tắt tài liệu nghiệp vụ dài và soạn sẵn email phản hồi."],
            ["Kỹ sư SRE / DevOps\n(SysAdmin / Lead Eng)", "Quản trị Hệ thống\n(Infrastructure Ops)", "Hệ thống sập, log lỗi tràn ngập, áp lực tìm nguyên nhân gốc rễ (RCA) nhanh.", "Công cụ bóc tách log tự động, khoanh vùng chính xác module lỗi và đề xuất Workaround."],
            ["Chuyên viên Pháp chế\n(Legal / Compliance)", "Giám sát Tuân thủ\n(Legal Counsel)", "Rà soát thủ công các điều khoản SLA, Uptime, DPA, GDPR trong hợp đồng nhà cung cấp.", "Trích xuất tự động Uptime %, MTTR, phát hiện các rủi ro thiếu bồi thường và vi phạm bảo mật."],
            ["Giám đốc Công nghệ\n(CIO / IT Director)", "Lãnh đạo Chiến lược\n(Executive Sponsor)", "Chi phí IT cao, rủi ro lộ lọt dữ liệu, thiếu báo cáo hiệu quả vận hành.", "Hệ thống bảo mật Zero-Trust, tối ưu chi phí Serverless và báo cáo minh bạch."],
        ]
    )

    # --- MỤC 3: PHẠM VI SẢN PHẨM & RÀNG BUỘC NGHIỆP VỤ ---
    add_styled_heading(doc, "3. PHẠM VI SẢN PHẨM & RÀNG BUỘC NGHIỆP VỤ (SCOPE & CONSTRAINTS)", 1)
    
    add_body_paragraph(
        doc,
        "• Phạm vi Trong Sản phẩm (In-Scope):\n"
        "  - Tiếp nhận và phân luồng tự động yêu cầu IT Helpdesk qua hội thoại tự nhiên.\n"
        "  - Hỗ trợ tự phục vụ: hướng dẫn reset mật khẩu, mở khóa tài khoản, thiết lập 2FA/MFA, VPN, Wi-Fi và máy in.\n"
        "  - Tra cứu tài liệu nghiệp vụ doanh nghiệp (ERP, HRM, CRM) qua giao thức MCP và BigQuery Vector Search.\n"
        "  - Phân tích log lỗi hệ thống (Root Cause Analysis) và rà soát SLA hợp đồng IT.\n"
        "  - Xác thực tập trung Google OIDC với cơ chế Fail-Closed Domain Filtering và phân quyền RBAC.\n"
        "  - Quản lý vòng đời ticket tích hợp Cloud Firestore.\n"
        "• Phạm vi Ngoài Sản phẩm (Out-of-Scope):\n"
        "  - Thay thế hoàn toàn quyết định phê duyệt mua sắm phần cứng hoặc chi ngân sách IT.\n"
        "  - Tự động thực thi lệnh reset tài khoản trên Domain Controller mà không có sự xác nhận của người dùng."
    )

    # --- MỤC 4: ĐẶC TẢ YÊU CẦU CHỨC NĂNG (FUNCTIONAL REQUIREMENTS) ---
    add_styled_heading(doc, "4. ĐẶC TẢ YÊU CẦU CHỨC NĂNG (FUNCTIONAL REQUIREMENTS - FRS)", 1)

    add_styled_heading(doc, "FR-01: Tầng Tiếp Nhận & Điều Phối Đa Tác Tử (Root Triage Orchestration)", 2)
    add_body_paragraph(
        doc,
        "• Mã chức năng: FR-01-TRIAGE\n"
        "• Tác tử thực thi: `root_triage_orchestrator` (Gemini 3 Flash Preview).\n"
        "• Yêu cầu chi tiết:\n"
        "  1. Tiếp nhận hội thoại từ người dùng, nạp lịch sử tương tác và ngữ cảnh thiết bị từ Vertex AI Memory Bank.\n"
        "  2. Tự động nhận diện ý định (Intent Recognition) và phân loại định tuyến:\n"
        "     - Định tuyến đến L1: Câu hỏi FAQ, quy trình tự phục vụ, yêu cầu tạo ticket sơ bộ.\n"
        "     - Định tuyến đến L2: Sự cố nghiệp vụ hệ thống ERP (SAP/Oracle), HRM (Workday), CRM (Salesforce).\n"
        "     - Định tuyến đến L3: Yêu cầu phân tích log lỗi hệ thống (RCA) hoặc rà soát SLA/DPA hợp đồng IT.\n"
        "  3. Tổng hợp kết quả phản hồi từ các Sub-agent và trả về cho người dùng với văn phong chuyên nghiệp."
    )

    add_styled_heading(doc, "FR-02: Mức 1 - Tự Phục Vụ & Xử Lý Sự Cố Cơ Bản (L1 Self-Service Specialist)", 2)
    add_body_paragraph(
        doc,
        "• Mã chức năng: FR-02-L1-SUPPORT\n"
        "• Tác tử thực thi: `l1_selfservice_agent` (Gemini 3 Flash Preview).\n"
        "• Yêu cầu chi tiết:\n"
        "  1. Hướng dẫn chi tiết từng bước (Step-by-step) các quy trình Self-Service: Reset mật khẩu Active Directory/Google Workspace/Okta, mở khóa tài khoản, kết nối mạng nội bộ.\n"
        "  2. Quản lý Ticket tự động: Sử dụng công cụ `create_helpdesk_ticket` tạo ticket mới với ID duy nhất (`TICK-XXXXXXXX`), thiết lập độ ưu tiên (`Low`, `Medium`, `High`, `Critical`) và ghi nhận vào Firestore.\n"
        "  3. Cung cấp chức năng tra cứu lịch sử ticket cá nhân (`list_user_tickets`) và chi tiết ticket (`get_ticket_details`)."
    )

    add_styled_heading(doc, "FR-03: Mức 2 - Tra Cứu Tri Thức Nghiệp Vụ Doanh Nghiệp (L2 Enterprise RAG)", 2)
    add_body_paragraph(
        doc,
        "• Mã chức năng: FR-03-L2-RAG\n"
        "• Tác tử thực thi: `l2_enterprise_rag_agent` (Gemini 3 Flash Preview) kết hợp FastMCP Server.\n"
        "• Yêu cầu chi tiết:\n"
        "  1. Tra cứu sâu cơ sở tri thức nghiệp vụ qua MCP Tools (`search_enterprise_knowledge`, `get_system_manual`):\n"
        "     - ERP: Lỗi phân quyền Purchase Order (PO), kỳ kế toán bị khóa, đồng bộ tồn kho.\n"
        "     - HRM: Lỗi đồng bộ máy chấm công vân tay, khóa kỳ tính lương Payroll, quy trình Onboarding.\n"
        "     - CRM: Lỗi đồng bộ Lead, vượt hạn ngạch API Limits, chuyển giao Account.\n"
        "  2. Tóm tắt tài liệu kỹ thuật dài (`summarize_long_document`): Trích xuất các điểm mấu chốt và hành động cần làm (Action Items).\n"
        "  3. Soạn thảo Email (`draft_email_response`): Soạn bản thảo email phản hồi khách hàng theo chuẩn mực doanh nghiệp."
    )

    add_styled_heading(doc, "FR-04: Mức 3 - Phân Tích Nguyên Nhân Gốc Rễ & Pháp Lý IT (L3 Deep Diagnostics)", 2)
    add_body_paragraph(
        doc,
        "• Mã chức năng: FR-04-L3-DIAGNOSTICS\n"
        "• Tác tử thực thi: `l3_deep_diagnostics_agent` (Gemini 3 Pro Preview - High Reasoning Model).\n"
        "• Yêu cầu chi tiết:\n"
        "  1. Root Cause Analysis (`analyze_system_logs_for_rca`): Phân tích log file, stack trace, phát hiện 6 nhóm dị thường cốt lõi (OOM, DB Connection Exhausted, Network Timeout, Auth Security Failure, Data Corruption Null, Disk I/O Failure) và lập báo cáo RCA chuẩn 4 phần.\n"
        "  2. Rà soát SLA Hợp đồng IT (`review_it_contract_sla`): Bóc tách cam kết Uptime %, thời gian phản hồi MTTR 2 chiều, phát hiện rủi ro thiếu điều khoản bồi thường (Service Credits), quyền kiểm toán và nghĩa vụ thông báo sự cố rò rỉ dữ liệu (24h–72h).\n"
        "  3. Kiểm soát quyền RBAC: Chỉ cho phép các vai trò đặc quyền (`it_admin`, `sys_admin`, `devops_engineer`, `compliance_officer`) thực thi."
    )

    add_styled_heading(doc, "FR-05: Lớp Bộ Đệm Ngữ Nghĩa (Semantic Cache Engine)", 2)
    add_body_paragraph(
        doc,
        "• Mã chức năng: FR-05-SEMANTIC-CACHE\n"
        "• Yêu cầu chi tiết:\n"
        "  1. Đánh giá độ tương đồng ngữ nghĩa câu hỏi bằng Vector Cosine Similarity (ngưỡng mặc định $\\ge 0.92$).\n"
        "  2. Khi Cache Hit: Trả lời ngay tức thì ($< 50\\text{ ms}$), không tiêu tốn token gọi Gemini.\n"
        "  3. Quản lý vòng đời bộ đệm: Tự động hết hạn sau 24h (TTL) và thu hồi theo chính sách LRU (Least Recently Used).\n"
        "  4. Cung cấp API giám sát: `GET /api/cache/stats` và `GET /api/cache/query`."
    )

    add_styled_heading(doc, "FR-06: Bảo Mật Định Danh & Phân Quyền Doanh Nghiệp (Enterprise SSO & RBAC)", 2)
    add_body_paragraph(
        doc,
        "• Mã chức năng: FR-06-SSO-RBAC\n"
        "• Yêu cầu chi tiết:\n"
        "  1. Xác thực Google Workspace OIDC ID Token chuẩn mực qua JWKS Public Keys.\n"
        "  2. Kiểm soát domain Fail-Closed (`ALLOWED_DOMAINS`): Chặn triệt để email ngoài tổ chức và Gmail cá nhân.\n"
        "  3. Chống tấn công Algorithm Confusion: Cô lập hoàn toàn giữa RS256 (OIDC Prod) và HS256 (Dev Mock Token).\n"
        "  4. Truyền tải ngữ cảnh người dùng qua `ContextVar` để kiểm tra phân quyền RBAC tại từng tool nhạy cảm."
    )

    # --- MỤC 5: ĐẶC TẢ YÊU CẦU PHI CHỨC NĂNG (NFRS) ---
    add_styled_heading(doc, "5. ĐẶC TẢ YÊU CẦU PHI CHỨC NĂNG (NON-FUNCTIONAL REQUIREMENTS - NFRS)", 1)

    nfr_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        nfr_table,
        [1.5, 1.5, 1.7, 1.8],
        ["Nhóm Tiêu chuẩn", "Chỉ số Mục tiêu", "Tiêu chí Chấp thuận (Acceptance Criteria)", "Phương pháp Đo lường"],
        [
            ["Hiệu năng (Performance)", "Cache Hit: < 50ms\nLLM Call: < 2.5s", "95% câu hỏi trong cache phản hồi < 50ms; 90% lượt gọi LLM phản hồi < 2.5s.", "Cloud Trace & Fast-API APM Metrics."],
            ["Độ sẵn sàng (Availability)", ">= 99.9% Uptime", "Không có điểm lỗi đơn (No Single Point of Failure). Cloud Run tự động khởi tạo lại container lỗi.", "Cloud Monitoring Uptime Checks."],
            ["Khả năng mở rộng (Scalability)", "0 -> 1,000 req/s", "Tự động scale out trên Cloud Run mà không bị nghẽn CPU hoặc cạn kiệt Connection Pool.", "Locust Load Testing Suite."],
            ["An toàn Thông tin (Security)", "Zero Trust & Least Privilege", "100% endpoint được bảo vệ bởi SSO Middleware. Quyền IAM Service Account được phân bổ tối thiểu.", "Báo cáo Kiểm thử Thâm nhập & SAIF Audit."],
            ["Khả năng phục hồi (Resilience)", "Graceful Degradation", "Khi Firestore hoặc BigQuery gián đoạn, hệ thống tự động fallback sang In-Memory mà không sập.", "Chaos Engineering / Unit Test Fallbacks."],
        ]
    )

    # --- MỤC 6: HÀNH TRÌNH NGƯỜI DÙNG & USE CASES TIÊU BIỂU ---
    add_styled_heading(doc, "6. HÀNH TRÌNH NGƯỜI DÙNG & KỊCH BẢN NGHIỆP VỤ (USE CASES)", 1)

    add_styled_heading(doc, "Use Case 01: Nhân viên yêu cầu Reset Mật khẩu Active Directory (L1)", 3)
    add_body_paragraph(
        doc,
        "1. Diễn viên: Nhân viên kinh doanh (Role: `employee`).\n"
        "2. Luồng sự kiện:\n"
        "   - Bước 1: Nhân viên gửi tin nhắn: 'Tôi bị khóa tài khoản máy tính do gõ sai pass nhiều lần, hỗ trợ reset giúp tôi'.\n"
        "   - Bước 2: Middleware xác thực token OIDC hợp lệ thuộc domain công ty.\n"
        "   - Bước 3: Orchestrator nhận diện ý định và định tuyến đến `l1_selfservice_agent`.\n"
        "   - Bước 4: L1 Agent hướng dẫn quy trình tự mở khóa qua cổng Self-Service Portal, đồng thời gọi `create_helpdesk_ticket` tạo ticket `TICK-AD-XXXX`.\n"
        "   - Bước 5: Phản hồi hướng dẫn chi tiết và mã ticket cho nhân viên."
    )

    add_styled_heading(doc, "Use Case 02: Kế toán viên báo lỗi Phân quyền Purchase Order trên SAP ERP (L2)", 3)
    add_body_paragraph(
        doc,
        "1. Diễn viên: Kế toán viên (Role: `employee`).\n"
        "2. Luồng sự kiện:\n"
        "   - Bước 1: Kế toán viên gửi: 'Tôi không thể duyệt đơn hàng PO-9981 trên SAP, hệ thống báo lỗi Authorization Check Failure M_EINK_FRG'.\n"
        "   - Bước 2: Orchestrator nhận diện từ khóa 'SAP', 'PO', 'Authorization' và chuyển tiếp cho `l2_enterprise_rag_agent`.\n"
        "   - Bước 3: L2 Agent gọi `search_enterprise_knowledge` từ MCP Server để tra cứu cẩm nang xử lý lỗi ERP.\n"
        "   - Bước 4: L2 Agent trích xuất giải pháp: cần cấp quyền qua transaction SU53 hoặc gửi yêu cầu phê duyệt cho Trưởng bộ phận Mua hàng, sau đó gọi `draft_email_response` soạn sẵn email mẫu gửi cho cấp quản lý."
    )

    add_styled_heading(doc, "Use Case 03: Kỹ sư DevOps gửi Log Trace phân tích sự cố sập máy chủ (L3)", 3)
    add_body_paragraph(
        doc,
        "1. Diễn viên: Kỹ sư DevOps (Role: `devops_engineer`).\n"
        "2. Luồng sự kiện:\n"
        "   - Bước 1: Kỹ sư gửi đoạn log stack trace máy chủ chứa các dòng 'java.lang.OutOfMemoryError: Java heap space' và 'exit code 137'.\n"
        "   - Bước 2: Orchestrator nhận diện log kỹ thuật và chuyển tiếp cho `l3_deep_diagnostics_agent` (Gemini 3 Pro Preview).\n"
        "   - Bước 3: L3 Agent gọi `analyze_system_logs_for_rca`. Hàm kiểm tra quyền RBAC: phát hiện user có role `devops_engineer` -> Cho phép thực thi.\n"
        "   - Bước 4: Động cơ phân tích phát hiện mẫu `OUT_OF_MEMORY` và sinh báo cáo RCA chuẩn 4 phần (Hiện tượng, Nguyên nhân gốc rễ rò rỉ bộ nhớ, Workaround nâng RAM và Kế hoạch phòng ngừa Heap Dump Profiling)."
    )

    add_styled_heading(doc, "Use Case 04: Chuyên viên Pháp chế Rà soát Cam kết SLA Hợp đồng Đám mây (L3)", 3)
    add_body_paragraph(
        doc,
        "1. Diễn viên: Chuyên viên Pháp chế (Role: `compliance_officer`).\n"
        "2. Luồng sự kiện:\n"
        "   - Bước 1: Chuyên viên gửi hợp đồng dịch vụ IT yêu cầu kiểm tra chỉ số SLA và rủi ro pháp lý.\n"
        "   - Bước 2: `l3_deep_diagnostics_agent` gọi `review_it_contract_sla`.\n"
        "   - Bước 3: Động cơ Regex bóc tách thành công: '99.95% Uptime', 'Thời gian phản hồi trong vòng 30 phút'.\n"
        "   - Bước 4: Hệ thống cảnh báo rủi ro pháp lý cao do hợp đồng thiếu điều khoản bồi thường thiệt hại tài chính (Service Credits) và thiếu cam kết thông báo sự cố rò rỉ dữ liệu trong vòng 72 giờ (DPA)."
    )

    # --- MỤC 7: MÔ HÌNH DỮ LIỆU & HỢP ĐỒNG GIAO TIẾP ---
    add_styled_heading(doc, "7. MÔ HÌNH DỮ LIỆU & HỢP ĐỒNG GIAO TIẾP (DATA CONTRACTS)", 1)

    add_body_paragraph(
        doc,
        "Cấu trúc dữ liệu chính thức của hệ thống được chuẩn hóa bằng Pydantic Models:"
    )

    add_code_block(
        doc,
        "# Schema Ticket Lưu trữ Đa tầng (Firestore / In-Memory)\n"
        "class HelpdeskTicket(BaseModel):\n"
        "    id: str = Field(description=\"Mã định danh duy nhất TICK-XXXXXXXX\")\n"
        "    user_id: str = Field(description=\"ID nhân viên yêu cầu\")\n"
        "    title: str = Field(description=\"Tiêu đề tóm tắt sự cố\")\n"
        "    description: str = Field(description=\"Mô tả chi tiết lỗi gặp phải\")\n"
        "    category: str = Field(default=\"General\", description=\"Phân loại sự cố IT\")\n"
        "    priority: Literal[\"Low\", \"Medium\", \"High\", \"Critical\"]\n"
        "    status: Literal[\"Open\", \"In_Progress\", \"Escalated\", \"Resolved\", \"Closed\"] = \"Open\"\n"
        "    assigned_tier: Literal[\"L1_SelfService\", \"L2_Enterprise_RAG\", \"L3_Deep_Diagnostics\", \"Human_Ops\"]\n"
        "    resolution_notes: Optional[str] = None\n"
        "    created_at: str\n"
        "    updated_at: str\n\n"
        "# Schema Người dùng Xác thực OIDC (SSOUser)\n"
        "class SSOUser(BaseModel):\n"
        "    user_id: str\n"
        "    email: str\n"
        "    email_verified: bool = True\n"
        "    full_name: str = \"Employee\"\n"
        "    department: str = \"General\"\n"
        "    roles: list[str] = [\"employee\"]\n"
        "    hosted_domain: Optional[str] = None\n"
        "    is_authenticated: bool = True"
    )

    # --- MỤC 8: LỘ TRÌNH PHÁT TRIỂN SẢN PHẨM ---
    add_styled_heading(doc, "8. LỘ TRÌNH PHÁT TRIỂN & KẾ HOẠCH BÀN GIAO (PRODUCT ROADMAP)", 1)

    roadmap_table = doc.add_table(rows=1, cols=4)
    format_custom_table(
        roadmap_table,
        [1.5, 1.5, 1.8, 1.7],
        ["Giai đoạn (Phase)", "Mốc Thời gian", "Mục tiêu Tính năng Chính", "Trạng thái Hiện tại"],
        [
            ["Phase 1: MVP Core", "Q1 / 2026", "Phân cấp 3-Tier Multi-Agent, Gemini 3 integration, in-memory knowledge store.", "ĐÃ HOÀN THÀNH (100%)"],
            ["Phase 2: Enterprise Ready", "Q2 / 2026", "Google OIDC SSO, Fail-Closed domain, BigQuery Vector Search, Semantic Cache, Terraform IaC.", "ĐÃ HOÀN THÀNH (100%)"],
            ["Phase 3: ITSM Integration", "Q3 / 2026", "Đồng bộ 2 chiều với ServiceNow, Jira Service Management, Slack/Teams Bot Connectors.", "Kế hoạch Quý 3"],
            ["Phase 4: Multimodal Voice", "Q4 / 2026", "Hỗ trợ cuộc gọi thoại trực tiếp qua Gemini Live Audio Streaming, OCR chụp màn hình lỗi.", "Kế hoạch Quý 4"],
        ]
    )

    add_body_paragraph(doc, "")
    add_body_paragraph(
        doc,
        "Tài liệu Đặc tả Yêu cầu Sản phẩm & Nghiệp vụ (SRS) này đã được hoàn thiện, kiểm chứng "
        "thông qua bộ kiểm thử tự động 46 test cases và sẵn sàng phục vụ cho quá trình nghiệm thu, bàn giao sản phẩm.",
        bold_prefix="KẾT LUẬN & PHÊ DUYỆT SẢN PHẨM: "
    )

    doc.save(output_path)
    print(f"Successfully generated Product SRS Document: {output_path}")

if __name__ == "__main__":
    target_path = "/Users/luuduc/.gemini/antigravity/scratch/it-helpdesk-agent/PRODUCT_SRS_DOCUMENT.docx"
    build_srs_docx(target_path)
