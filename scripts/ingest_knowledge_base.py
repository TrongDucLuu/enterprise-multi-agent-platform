#!/usr/bin/env python3
"""
Enterprise Knowledge Base Data Ingestion Pipeline.

Parses multi-format customer documents (.md, .txt, .docx, .pdf, .jsonl), validates enterprise system
tags against systems.yaml config, splits documents into semantic chunks, generates 768-dimensional
dense vector embeddings using text-embedding-005, and performs idempotent upsert (MERGE) with
orphaned chunks cleanup via temporary Staging Table into BigQuery.

Usage:
    # Dry-run parsing and embedding simulation:
    python scripts/ingest_knowledge_base.py --source-dir data/knowledge_base/ --dry-run

    # Production ingestion into BigQuery:
    python scripts/ingest_knowledge_base.py --source-dir data/knowledge_base/ --project-id my-project --dataset-id it_helpdesk_kb

    # Ingest single document and run test query:
    python scripts/ingest_knowledge_base.py --file docs/sap_procurement_guide.pdf --system ERP --test-query "lỗi phân quyền ME21N"
"""

import os
import sys
import re
import time
import json
import uuid
import hashlib
import argparse
import logging
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Optional, Any

# Ensure project root is on sys.path
BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from it_helpdesk_agent.app_utils.system_config import (
    get_configured_systems,
    get_valid_system_filters,
    get_system_metadata,
    get_chunking_config,
    get_document_processing_config,
)
from it_helpdesk_agent.app_utils.embedding_utils import (
    DEFAULT_EMBEDDING_MODEL,
    generate_batch_embeddings,
    generate_text_embedding,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger("ingest_knowledge_base")


class DocumentParser:
    """Extracts raw text, structured sections, and metadata from various enterprise document formats."""

    @staticmethod
    def parse_markdown_or_text(file_path: Path) -> list[dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        title = file_path.stem.replace("_", " ").title()
        heading_pattern = re.compile(r"^(#{1,3})\s+(.+)$")
        
        sections: list[dict[str, Any]] = []
        current_heading = title
        current_level = 1
        current_h1: Optional[str] = title
        current_h2: Optional[str] = None
        current_h3: Optional[str] = None
        current_lines: list[str] = []

        for line in content.splitlines():
            m = heading_pattern.match(line)
            if m:
                if current_lines:
                    sec_text = "\n".join(current_lines).strip()
                    if sec_text:
                        sections.append({
                            "level": current_level,
                            "heading": current_heading,
                            "content": sec_text,
                            "hierarchy": {
                                "h1": current_h1,
                                "h2": current_h2,
                                "h3": current_h3,
                            },
                        })
                    current_lines = []
                current_level = len(m.group(1))
                current_heading = m.group(2).strip()
                if current_level == 1:
                    current_h1 = current_heading
                    current_h2 = None
                    current_h3 = None
                    if not sections:
                        title = current_heading
                elif current_level == 2:
                    current_h2 = current_heading
                    current_h3 = None
                elif current_level == 3:
                    current_h3 = current_heading
            else:
                current_lines.append(line)

        if current_lines:
            sec_text = "\n".join(current_lines).strip()
            if sec_text:
                sections.append({
                    "level": current_level,
                    "heading": current_heading,
                    "content": sec_text,
                    "hierarchy": {
                        "h1": current_h1,
                        "h2": current_h2,
                        "h3": current_h3,
                    },
                })

        return [{
            "title": title,
            "content": content,
            "source_uri": str(file_path),
            "file_type": file_path.suffix.lower(),
            "sections": sections,
        }]

    @staticmethod
    def parse_docx(file_path: Path) -> list[dict[str, Any]]:
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            title = paragraphs[0] if paragraphs else file_path.stem.replace("_", " ").title()
            
            sections: list[dict[str, Any]] = []
            current_heading = title
            current_level = 1
            current_h1: Optional[str] = title
            current_h2: Optional[str] = None
            current_h3: Optional[str] = None
            current_paras: list[str] = []

            for p in doc.paragraphs:
                p_text = p.text.strip()
                if not p_text:
                    continue
                style_name = getattr(p.style, "name", "") or ""
                if "heading" in style_name.lower():
                    if current_paras:
                        sec_text = "\n\n".join(current_paras).strip()
                        if sec_text:
                            sections.append({
                                "level": current_level,
                                "heading": current_heading,
                                "content": sec_text,
                                "hierarchy": {
                                    "h1": current_h1,
                                    "h2": current_h2,
                                    "h3": current_h3,
                                },
                            })
                        current_paras = []
                    m = re.search(r"\d+", style_name)
                    current_level = int(m.group(0)) if m else 1
                    current_heading = p_text
                    if current_level == 1:
                        current_h1 = current_heading
                        current_h2 = None
                        current_h3 = None
                        if not sections:
                            title = current_heading
                    elif current_level == 2:
                        current_h2 = current_heading
                        current_h3 = None
                    elif current_level == 3:
                        current_h3 = current_heading
                else:
                    current_paras.append(p_text)

            if current_paras:
                sec_text = "\n\n".join(current_paras).strip()
                if sec_text:
                    sections.append({
                        "level": current_level,
                        "heading": current_heading,
                        "content": sec_text,
                        "hierarchy": {
                            "h1": current_h1,
                            "h2": current_h2,
                            "h3": current_h3,
                        },
                    })

            return [{
                "title": title,
                "content": content,
                "source_uri": str(file_path),
                "file_type": ".docx",
                "sections": sections,
            }]
        except ImportError:
            logger.warning("python-docx is not installed. To parse DOCX files, install python-docx (`pip install python-docx`).")
            return []
        except Exception as e:
            logger.error("Failed to parse docx file %s: %s", file_path, e)
            return []

    @staticmethod
    def parse_pdf_document_ai(file_path: Path, config: dict[str, Any]) -> list[dict[str, Any]]:
        """
        Parses a PDF using Google Cloud Document AI Layout Parser.
        Extracts structured sections (heading-1, heading-2, paragraphs, tables).
        Retries up to max_retries with exponential backoff and timeout.
        Fails closed on unrecoverable errors.
        """
        processor_id = config.get("document_ai_processor_id")
        if not processor_id:
            raise RuntimeError("Missing 'document_ai_processor_id' for Document AI PDF parsing. (Fail-Closed)")

        timeout_seconds = config.get("document_ai_timeout_seconds", 60.0)
        max_retries = config.get("document_ai_max_retries", 2)

        try:
            from google.cloud import documentai
        except ImportError:
            logger.error("google-cloud-documentai is not installed. Run `pip install google-cloud-documentai`.")
            raise

        with open(file_path, "rb") as f:
            pdf_bytes = f.read()

        raw_document = documentai.RawDocument(content=pdf_bytes, mime_type="application/pdf")
        request = documentai.ProcessRequest(name=processor_id, raw_document=raw_document)

        client = documentai.DocumentProcessorServiceClient()

        last_error = None
        for attempt in range(max_retries + 1):
            try:
                logger.info(
                    "Calling Document AI Layout Parser for %s (attempt %d/%d, timeout=%.1fs)...",
                    file_path.name, attempt + 1, max_retries + 1, timeout_seconds
                )
                result = client.process_document(request=request, timeout=timeout_seconds)
                doc = result.document
                return DocumentParser._map_document_ai_to_document_info(doc, file_path)
            except Exception as e:
                last_error = e
                logger.warning("Document AI attempt %d failed for %s: %s", attempt + 1, file_path.name, e)
                if attempt < max_retries:
                    backoff = 2 ** attempt
                    time.sleep(backoff)

        raise RuntimeError(
            f"Document AI parsing failed for '{file_path}' after {max_retries + 1} attempts: {last_error}. (Fail-Closed)"
        )

    @staticmethod
    def _map_document_ai_to_document_info(doc: Any, file_path: Path) -> list[dict[str, Any]]:
        full_text = getattr(doc, "text", None) if not isinstance(doc, dict) else doc.get("text", "")
        if full_text is None:
            full_text = ""

        blocks = []
        pages = getattr(doc, "pages", []) if not isinstance(doc, dict) else doc.get("pages", [])
        for page in pages:
            page_blocks = getattr(page, "blocks", []) if not isinstance(page, dict) else page.get("blocks", [])
            for b in page_blocks:
                blocks.append(b)

        title = file_path.stem.replace("_", " ").title()
        sections: list[dict[str, Any]] = []
        current_heading = title
        current_level = 1
        current_h1: Optional[str] = title
        current_h2: Optional[str] = None
        current_h3: Optional[str] = None
        current_content_parts: list[str] = []

        def get_block_type_and_text(block: Any) -> tuple[str, str]:
            if isinstance(block, dict):
                b_type = block.get("type_") or block.get("block_type") or block.get("type", "paragraph")
                b_text = block.get("text", "")
                if not b_text and "layout" in block:
                    anchors = block["layout"].get("text_anchor", {}).get("text_segments", [])
                    b_text = "".join(full_text[int(seg.get("start_index", 0)):int(seg.get("end_index", 0))] for seg in anchors)
                return str(b_type), b_text.strip()
            else:
                b_type = getattr(block, "type_", None) or getattr(block, "block_type", "paragraph")
                b_text = getattr(block, "text", "")
                if not b_text and hasattr(block, "layout") and hasattr(block.layout, "text_anchor"):
                    segments = getattr(block.layout.text_anchor, "text_segments", [])
                    b_text = "".join(full_text[int(getattr(s, "start_index", 0)):int(getattr(s, "end_index", 0))] for s in segments)
                return str(b_type), b_text.strip()

        for b in blocks:
            b_type, b_text = get_block_type_and_text(b)
            if not b_text:
                continue

            if "heading" in b_type.lower():
                if current_content_parts:
                    sec_text = "\n\n".join(current_content_parts).strip()
                    if sec_text:
                        sections.append({
                            "level": current_level,
                            "heading": current_heading,
                            "content": sec_text,
                            "hierarchy": {
                                "h1": current_h1,
                                "h2": current_h2,
                                "h3": current_h3,
                            },
                        })
                    current_content_parts = []
                
                m = re.search(r"\d+", b_type)
                current_level = int(m.group(0)) if m else 1
                current_heading = b_text
                if current_level == 1:
                    current_h1 = current_heading
                    current_h2 = None
                    current_h3 = None
                    if not sections:
                        title = current_heading
                elif current_level == 2:
                    current_h2 = current_heading
                    current_h3 = None
                elif current_level == 3:
                    current_h3 = current_heading
            else:
                current_content_parts.append(b_text)

        if current_content_parts:
            sec_text = "\n\n".join(current_content_parts).strip()
            if sec_text:
                sections.append({
                    "level": current_level,
                    "heading": current_heading,
                    "content": sec_text,
                    "hierarchy": {
                        "h1": current_h1,
                        "h2": current_h2,
                        "h3": current_h3,
                    },
                })

        return [{
            "title": title,
            "content": full_text.strip() if full_text.strip() else "\n\n".join(s["content"] for s in sections),
            "source_uri": str(file_path),
            "file_type": ".pdf",
            "sections": sections,
        }]


    @staticmethod
    def parse_pdf(file_path: Path, doc_proc_config: Optional[dict[str, Any]] = None) -> list[dict[str, Any]]:
        if doc_proc_config is None:
            try:
                doc_proc_config = get_document_processing_config()
            except Exception as e:
                logger.debug("Could not load document processing config: %s", e)
                doc_proc_config = {"pdf_parser": "pypdf_flat"}

        pdf_parser_mode = (doc_proc_config or {}).get("pdf_parser", "pypdf_flat")
        if pdf_parser_mode == "document_ai":
            return DocumentParser.parse_pdf_document_ai(file_path, doc_proc_config)

        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text.strip())
            content = "\n\n".join(pages_text)
            title = file_path.stem.replace("_", " ").title()
            if pages_text:
                first_lines = [l.strip() for l in pages_text[0].splitlines() if l.strip()]
                if first_lines:
                    title = first_lines[0][:100]
            return [{
                "title": title,
                "content": content,
                "source_uri": str(file_path),
                "file_type": ".pdf",
                "sections": [],
            }]
        except ImportError:
            logger.warning("pypdf is not installed. To parse PDF files, install pypdf (`pip install pypdf`).")
            return []
        except Exception as e:
            logger.error("Failed to parse pdf file %s: %s", file_path, e)
            return []

    @staticmethod
    def parse_jsonl(file_path: Path) -> list[dict[str, Any]]:
        articles = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line_no, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    # Gán line_no vào default source_uri để phân biệt các dòng khác nhau trong cùng file JSONL
                    source_uri = data.get("source_uri") or f"{file_path}#L{line_no}"
                    articles.append({
                        "id": data.get("id"),
                        "system": data.get("system"),
                        "title": data.get("title", f"Article {line_no}"),
                        "category": data.get("category", "General"),
                        "content": data.get("content", ""),
                        "keywords": data.get("keywords", []),
                        "source_uri": source_uri,
                        "file_type": ".jsonl",
                        "sections": [],
                    })
                except json.JSONDecodeError as e:
                    logger.warning("Invalid JSON at line %d in %s: %s", line_no, file_path, e)
        return articles


def is_well_structured(
    sections: Optional[list[dict[str, Any]]],
    max_chunk_size: int = 1200,
    max_section_ratio: float = 0.65,
    min_avg_length: int = 100,
) -> bool:
    """
    Evaluates whether a document's extracted sections provide clean structural boundaries:
    - At least 2 sections.
    - No single section occupies more than max_section_ratio (65%) of total document text.
    - Average length per section is at least min_avg_length (100 chars).
    """
    if not sections or len(sections) < 2:
        return False

    total_len = sum(len(s.get("content", "")) for s in sections)
    if total_len == 0:
        return False

    avg_len = total_len / len(sections)
    if avg_len < min_avg_length:
        return False

    for s in sections:
        sec_len = len(s.get("content", ""))
        if (sec_len / total_len) > max_section_ratio:
            return False

    return True


def chunk_by_sections(
    sections: list[dict[str, Any]],
    max_chunk_size: int = 1200,
    overlap: int = 150,
    return_metadata: bool = False,
) -> Any:
    """
    Chunks document section-by-section, keeping headings attached to content.
    If a section exceeds max_chunk_size, it is recursively split within that section's scope.
    """
    chunks = []
    chunk_meta = []
    for sec in sections:
        heading = sec.get("heading", "").strip()
        content = sec.get("content", "").strip()
        hierarchy = sec.get("hierarchy") or {
            "h1": heading or None,
            "h2": None,
            "h3": None,
        }
        if not content:
            continue

        header_prefix = f"## {heading}\n\n" if heading else ""
        full_sec_text = f"{header_prefix}{content}".strip()

        if len(full_sec_text) <= max_chunk_size:
            chunks.append(full_sec_text)
            chunk_meta.append(hierarchy)
        else:
            # Section exceeds max_chunk_size -> recursive split content within section scope
            sub_max_size = max(200, max_chunk_size - len(header_prefix))
            sub_chunks = chunk_text(content, max_chunk_size=sub_max_size, overlap=overlap)
            for sub in sub_chunks:
                chunks.append(f"{header_prefix}{sub}".strip())
                chunk_meta.append(hierarchy)

    if return_metadata:
        return [{"text": c, "hierarchy": m} for c, m in zip(chunks, chunk_meta) if c]
    return [c for c in chunks if c]


def chunk_text(
    text: str,
    max_chunk_size: int = 1200,
    overlap: int = 150,
    separators: Optional[list[str]] = None,
) -> list[str]:
    """
    Splits text recursively using prioritized separators:
    \n\n\n -> \n\n -> \n -> .  -> hard character split.
    """
    if not text or not text.strip():
        return []

    text = text.strip()
    if len(text) <= max_chunk_size:
        return [text]

    if separators is None:
        separators = ["\n\n\n", "\n\n", "\n", ". "]

    chosen_sep = None
    for sep in separators:
        if sep in text:
            chosen_sep = sep
            break

    if chosen_sep is None:
        # Fallback to hard character split
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + max_chunk_size, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
        return chunks

    splits = text.split(chosen_sep)
    chunks = []
    current_parts: list[str] = []
    current_len = 0
    remaining_seps = separators[separators.index(chosen_sep) + 1:]

    for part in splits:
        if not part.strip():
            continue
        part_len = len(part)

        if part_len > max_chunk_size:
            if current_parts:
                merged = chosen_sep.join(current_parts).strip()
                if merged:
                    chunks.append(merged)
                current_parts = []
                current_len = 0
            sub_chunks = chunk_text(part, max_chunk_size=max_chunk_size, overlap=overlap, separators=remaining_seps)
            chunks.extend(sub_chunks)
        elif current_len + (len(chosen_sep) if current_parts else 0) + part_len <= max_chunk_size:
            current_parts.append(part)
            current_len += (len(chosen_sep) if len(current_parts) > 1 else 0) + part_len
        else:
            merged = chosen_sep.join(current_parts).strip()
            if merged:
                chunks.append(merged)

            overlap_parts: list[str] = []
            overlap_len = 0
            for p in reversed(current_parts):
                if overlap_len + len(p) <= overlap:
                    overlap_parts.insert(0, p)
                    overlap_len += len(p) + len(chosen_sep)
                else:
                    break
            current_parts = overlap_parts + [part]
            current_len = sum(len(p) for p in current_parts) + (len(chosen_sep) * max(0, len(current_parts) - 1))

    if current_parts:
        merged = chosen_sep.join(current_parts).strip()
        if merged:
            chunks.append(merged)

    return chunks


def process_document(
    doc_info: dict[str, Any],
    default_system: Optional[str] = None
) -> list[dict[str, Any]]:
    """
    Validates, chunks, and prepares a document for embedding and ingestion.
    Applies customer-configurable tiered chunking strategy (auto, fixed, semantic),
    generates deterministic article IDs based on SHA-256(system:source_uri:title:idx)
    and computes content_hash for CDC change tracking.
    """
    valid_systems = get_valid_system_filters()
    configured_systems = get_configured_systems()

    raw_system = doc_info.get("system") or default_system
    if not raw_system:
        # Try inferring system from title/filename
        text_for_infer = (doc_info.get("title", "") + " " + doc_info.get("source_uri", "")).upper()
        for s in configured_systems:
            if s in text_for_infer:
                raw_system = s
                break

    if not raw_system:
        raw_system = configured_systems[0] if configured_systems else "ERP"
        logger.info("Defaulting system to '%s' for '%s'", raw_system, doc_info.get("title"))

    system_clean = raw_system.strip().upper()
    if system_clean not in valid_systems or system_clean == "ALL":
        raise ValueError(
            f"Hệ thống '{raw_system}' không hợp lệ hoặc là từ khóa dành riêng. "
            f"Các hệ thống được hỗ trợ: {configured_systems}"
        )

    title = doc_info.get("title", "Untitled Document")
    category = doc_info.get("category", "Operations")
    content = doc_info.get("content", "")
    source_uri = doc_info.get("source_uri", "")
    raw_keywords = doc_info.get("keywords", [])
    sections = doc_info.get("sections", [])

    # Load system-specific or global chunking configuration
    chunking_cfg = get_chunking_config(system_clean)
    strategy = chunking_cfg.get("strategy", "auto")
    max_chunk_size = chunking_cfg.get("max_chunk_size", 1200)
    overlap = chunking_cfg.get("overlap", 150)
    max_sec_ratio = chunking_cfg.get("well_structured_max_section_ratio", 0.65)
    min_avg_len = chunking_cfg.get("well_structured_min_avg_section_length", 100)

    chunk_items: list[dict[str, Any]] = []

    if strategy == "semantic":
        logger.info("Semantic chunking strategy flagged for system '%s' (fallback to structured/recursive)", system_clean)
        if sections and is_well_structured(sections, max_chunk_size=max_chunk_size, max_section_ratio=max_sec_ratio, min_avg_length=min_avg_len):
            chunk_items = chunk_by_sections(sections, max_chunk_size=max_chunk_size, overlap=overlap, return_metadata=True)
        else:
            raw_c = chunk_text(content, max_chunk_size=max_chunk_size, overlap=overlap)
            chunk_items = [{"text": c, "hierarchy": {"h1": title, "h2": None, "h3": None}} for c in raw_c]
    elif strategy == "fixed":
        raw_c = chunk_text(content, max_chunk_size=max_chunk_size, overlap=overlap)
        chunk_items = [{"text": c, "hierarchy": {"h1": title, "h2": None, "h3": None}} for c in raw_c]
    else:  # "auto"
        if sections and is_well_structured(sections, max_chunk_size=max_chunk_size, max_section_ratio=max_sec_ratio, min_avg_length=min_avg_len):
            chunk_items = chunk_by_sections(sections, max_chunk_size=max_chunk_size, overlap=overlap, return_metadata=True)
        else:
            raw_c = chunk_text(content, max_chunk_size=max_chunk_size, overlap=overlap)
            chunk_items = [{"text": c, "hierarchy": {"h1": title, "h2": None, "h3": None}} for c in raw_c]

    processed_articles = []

    for idx, item in enumerate(chunk_items):
        chunk = item["text"]
        section_hierarchy = item.get("hierarchy") or {"h1": title, "h2": None, "h3": None}

        # Generate deterministic ID
        if doc_info.get("id") and len(chunk_items) == 1:
            article_id = doc_info["id"].upper()
        else:
            hasher = hashlib.sha256()
            hasher.update(f"{system_clean}:{source_uri}:{title}:{idx}".encode("utf-8"))
            short_hash = hasher.hexdigest()[:8].upper()
            article_id = f"{system_clean}-KB-{short_hash}"

        chunk_title = title if len(chunk_items) == 1 else f"{title} (Phần {idx + 1}/{len(chunk_items)})"
        
        # Extract keywords if not provided
        keywords = raw_keywords.copy()
        if not keywords:
            words = [w.strip(".,;:()") for w in chunk_title.lower().split() if len(w) > 2]
            keywords = list(set(words))[:8]

        # Compute content_hash for Change Data Capture (CDC)
        content_hash = hashlib.sha256(chunk.encode("utf-8")).hexdigest()

        processed_articles.append({
            "id": article_id,
            "system": system_clean,
            "title": chunk_title,
            "category": category,
            "content": chunk,
            "keywords": keywords,
            "source_uri": source_uri,
            "content_hash": content_hash,
            "section_hierarchy": section_hierarchy,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        })

    return processed_articles


def ensure_vector_index(
    bq_client: Any,
    project_id: str,
    dataset_id: str,
    table_name: str = "knowledge_articles",
    index_name: str = "knowledge_articles_vector_idx"
):
    """
    Executes BigQuery CREATE VECTOR INDEX DDL if index does not exist with STORING clause.
    BigQuery IVF Vector Index will automatically optimize vector queries;
    if dataset has fewer than 5,000 rows, BigQuery will automatically use exact cosine search.
    """
    ddl = f"""
    CREATE VECTOR INDEX IF NOT EXISTS `{index_name}`
    ON `{project_id}.{dataset_id}.{table_name}`(embedding)
    STORING (system, category, id, title, content, section_hierarchy)
    OPTIONS(distance_type='COSINE', index_type='IVF')
    """
    try:
        logger.info("Verifying / Creating BigQuery Vector Index '%s' with STORING columns...", index_name)
        query_job = bq_client.query(ddl)
        query_job.result()
        logger.info("BigQuery Vector Index '%s' is verified and active.", index_name)
    except Exception as e:
        logger.warning(
            "Note: BigQuery Vector Index DDL returned: %s. "
            "(BigQuery automatically executes exact cosine search when dataset size is under 5,000 rows threshold).",
            e
        )


def check_vector_index_coverage(
    bq_client: Any,
    project_id: str,
    dataset_id: str,
    table_name: str = "knowledge_articles",
    index_name: str = "knowledge_articles_vector_idx"
) -> dict[str, Any]:
    """
    Monitors BigQuery Vector Index status, coverage percentage, and unindexed row count via INFORMATION_SCHEMA.
    Logs clear operational diagnostics for enterprise observability.
    """
    coverage_sql = f"""
    SELECT 
        table_name,
        index_name,
        index_status,
        coverage_percentage,
        unindexed_row_count,
        total_row_count
    FROM `{project_id}.{dataset_id}.INFORMATION_SCHEMA.VECTOR_INDEXES`
    WHERE table_name = @table_name AND index_name = @index_name
    """
    try:
        from google.cloud import bigquery
        job_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("table_name", "STRING", table_name),
                bigquery.ScalarQueryParameter("index_name", "STRING", index_name),
            ]
        )
        rows = list(bq_client.query(coverage_sql, job_config=job_config).result())
        if not rows:
            logger.info("Vector Index '%s' does not exist yet in INFORMATION_SCHEMA or is newly scheduled.", index_name)
            return {"index_status": "NOT_FOUND", "coverage_percentage": 0.0}

        row = rows[0]
        status = getattr(row, "index_status", "UNKNOWN")
        coverage = getattr(row, "coverage_percentage", 0.0) or 0.0
        unindexed = getattr(row, "unindexed_row_count", 0) or 0
        total = getattr(row, "total_row_count", 0) or 0

        if coverage == 0.0:
            if status == "TEMPORARILY DISABLED":
                logger.info(
                    "Lưu ý: Vector Index '%s' có trạng thái 'TEMPORARILY DISABLED' (Coverage: 0.0%%). "
                    "Nguyên nhân: Kích thước bảng tri thức dưới ngưỡng tối thiểu (thường < 10 MB) "
                    "nên BigQuery tự động dùng Exact Cosine Search. Đây là hành vi bình thường cho cơ sở tri thức nhỏ của khách hàng mới.",
                    index_name
                )
            else:
                logger.warning(
                    "CẢNH BÁO: Vector Index '%s' coverage = 0.0%% (Status: %s). "
                    "Truy vấn sẽ thực hiện Full Table Scan cho tới khi index hoàn tất indexing.",
                    index_name, status
                )
        else:
            logger.info(
                "Vector Index '%s' đang hoạt động tốt. Status: %s, Coverage: %.1f%%, Dòng chưa index: %d / %d tổng số dòng.",
                index_name, status, coverage, unindexed, total
            )

        return {
            "index_status": status,
            "coverage_percentage": coverage,
            "unindexed_row_count": unindexed,
            "total_row_count": total,
        }
    except Exception as e:
        logger.warning("Could not query INFORMATION_SCHEMA.VECTOR_INDEXES: %s", e)
        return {"index_status": "ERROR", "error": str(e)}


def ingest_articles_to_bigquery(
    articles: list[dict[str, Any]],
    project_id: str,
    dataset_id: str,
    table_name: str = "knowledge_articles"
) -> int:
    """
    Performs production-grade idempotent upsert (MERGE) into BigQuery:
    1. CDC pre-check on content_hash to skip redundant embedding API calls.
    2. Batch loads articles into a temporary staging table (zero streaming buffer locks on target).
    3. Executes atomic SQL MERGE from staging table into target table with section_hierarchy.
    4. Executes DML DELETE on target table to clean up orphaned chunks for modified documents.
    5. Drops staging table, ensures BigQuery IVF Vector Index with STORING is active, and monitors coverage.
    """
    if not articles:
        logger.info("No articles to ingest.")
        return 0

    # Deduplicate input articles by 'id', keeping the latest entry
    deduped_articles: dict[str, dict[str, Any]] = {}
    duplicate_id_count = 0
    duplicate_ids_sample: list[str] = []
    for a in articles:
        art_id = a.get("id")
        if art_id:
            if art_id in deduped_articles:
                duplicate_id_count += 1
                if len(duplicate_ids_sample) < 5:
                    duplicate_ids_sample.append(art_id)
            deduped_articles[art_id] = a
        else:
            deduped_articles[str(uuid.uuid4())] = a

    if duplicate_id_count > 0:
        logger.warning(
            "Phát hiện %d chunk trùng ID trong tập nạp (ví dụ các ID: %s). "
            "Hệ thống đã tự động loại bỏ bản ghi cũ và giữ bản ghi mới nhất để bảo vệ an toàn cho câu lệnh MERGE.",
            duplicate_id_count,
            ", ".join(duplicate_ids_sample)
        )
    articles = list(deduped_articles.values())

    try:
        from google.cloud import bigquery
    except ImportError:
        logger.error("google-cloud-bigquery is not installed. Please install dependencies.")
        raise

    bq_client = bigquery.Client(project=project_id)
    full_target_table = f"`{project_id}.{dataset_id}.{table_name}`"

    # 1. CDC Pre-Check: Retrieve existing content hashes to avoid redundant embedding generation
    source_uris = list({a["source_uri"] for a in articles if a.get("source_uri")})
    existing_hashes: dict[str, str] = {}
    existing_embeddings: dict[str, list[float]] = {}

    if source_uris:
        try:
            cdc_sql = f"""
            SELECT id, content_hash, embedding
            FROM {full_target_table}
            WHERE source_uri IN UNNEST(@source_uris)
            """
            job_config = bigquery.QueryJobConfig(
                query_parameters=[
                    bigquery.ArrayQueryParameter("source_uris", "STRING", source_uris)
                ]
            )
            cdc_rows = list(bq_client.query(cdc_sql, job_config=job_config).result())
            for r in cdc_rows:
                if hasattr(r, "id") and hasattr(r, "content_hash") and r.id and r.content_hash:
                    existing_hashes[r.id] = r.content_hash
                    if hasattr(r, "embedding") and r.embedding:
                        existing_embeddings[r.id] = list(r.embedding)
        except Exception as e:
            logger.debug("CDC pre-check bypassed (table may be newly initialized): %s", e)

    # 2. Selective Embedding Generation
    chunks_to_embed_indices: list[int] = []
    texts_to_embed: list[str] = []

    for idx, a in enumerate(articles):
        art_id = a.get("id")
        content_hash = a.get("content_hash")
        # Reuse existing vector if content is unchanged
        if art_id and art_id in existing_hashes and existing_hashes[art_id] == content_hash and art_id in existing_embeddings:
            a["embedding"] = existing_embeddings[art_id]
        elif a.get("embedding"):
            # Already has embedding (e.g. injected or dry-run)
            pass
        else:
            chunks_to_embed_indices.append(idx)
            texts_to_embed.append(f"{a['title']}\n{a['category']}\n{a['content']}")

    reused_count = len(articles) - len(chunks_to_embed_indices)
    if reused_count > 0:
        logger.info("CDC Optimization: Reused %d existing embeddings (content unchanged).", reused_count)

    if chunks_to_embed_indices:
        logger.info("Generating embeddings for %d new/modified chunks using %s...", len(chunks_to_embed_indices), DEFAULT_EMBEDDING_MODEL)
        new_embeddings = generate_batch_embeddings(texts_to_embed, model_name=DEFAULT_EMBEDDING_MODEL)
        for idx, emb in zip(chunks_to_embed_indices, new_embeddings):
            articles[idx]["embedding"] = emb

    # 3. Create Temporary Staging Table
    staging_suffix = uuid.uuid4().hex[:8]
    staging_table_name = f"{table_name}_staging_{staging_suffix}"
    staging_table_id = f"{project_id}.{dataset_id}.{staging_table_name}"
    full_staging_table = f"`{project_id}.{dataset_id}.{staging_table_name}`"

    schema = [
        bigquery.SchemaField("id", "STRING", mode="REQUIRED"),
        bigquery.SchemaField("system", "STRING", mode="REQUIRED"),
        bigquery.SchemaField("title", "STRING", mode="REQUIRED"),
        bigquery.SchemaField("category", "STRING", mode="NULLABLE"),
        bigquery.SchemaField("content", "STRING", mode="REQUIRED"),
        bigquery.SchemaField("keywords", "STRING", mode="REPEATED"),
        bigquery.SchemaField("embedding", "FLOAT64", mode="REPEATED"),
        bigquery.SchemaField("source_uri", "STRING", mode="NULLABLE"),
        bigquery.SchemaField("content_hash", "STRING", mode="NULLABLE"),
        bigquery.SchemaField(
            "section_hierarchy",
            "RECORD",
            mode="NULLABLE",
            fields=[
                bigquery.SchemaField("h1", "STRING", mode="NULLABLE"),
                bigquery.SchemaField("h2", "STRING", mode="NULLABLE"),
                bigquery.SchemaField("h3", "STRING", mode="NULLABLE"),
            ]
        ),
        bigquery.SchemaField("updated_at", "TIMESTAMP", mode="REQUIRED"),
    ]

    staging_table = bigquery.Table(staging_table_id, schema=schema)
    staging_table.expires = datetime.now(timezone.utc) + timedelta(hours=1)
    bq_client.create_table(staging_table, exists_ok=True)
    logger.info("Created temporary staging table %s", staging_table_name)

    try:
        # 4. Batch Load Articles into Staging Table (Load Job - Free of streaming buffer locks)
        load_job_config = bigquery.LoadJobConfig(
            schema=schema,
            write_disposition=bigquery.WriteDisposition.WRITE_TRUNCATE,
        )
        load_job = bq_client.load_table_from_json(articles, staging_table_id, job_config=load_job_config)
        load_job.result()
        logger.info("Loaded %d articles into staging table.", len(articles))

        # 5. Execute Atomic SQL MERGE from Staging into Target Table with deduplication fail-safe
        merge_sql = f"""
        MERGE {full_target_table} T
        USING (
          SELECT * FROM {full_staging_table}
          QUALIFY ROW_NUMBER() OVER (PARTITION BY id ORDER BY updated_at DESC) = 1
        ) S
        ON T.id = S.id
        WHEN MATCHED AND (T.content_hash != S.content_hash OR T.content_hash IS NULL) THEN
          UPDATE SET
            T.system = S.system,
            T.title = S.title,
            T.category = S.category,
            T.content = S.content,
            T.keywords = S.keywords,
            T.embedding = S.embedding,
            T.source_uri = S.source_uri,
            T.content_hash = S.content_hash,
            T.section_hierarchy = S.section_hierarchy,
            T.updated_at = S.updated_at
        WHEN NOT MATCHED THEN
          INSERT (id, system, title, category, content, keywords, embedding, source_uri, content_hash, section_hierarchy, updated_at)
          VALUES (S.id, S.system, S.title, S.category, S.content, S.keywords, S.embedding, S.source_uri, S.content_hash, S.section_hierarchy, S.updated_at);
        """
        logger.info("Executing Atomic MERGE into %s...", full_target_table)
        merge_job = bq_client.query(merge_sql)
        merge_job.result()
        logger.info("Atomic MERGE completed successfully.")

        # 6. Execute Orphaned Chunks Cleanup via DML on Target Table (No streaming buffer lock on target!)
        if source_uris:
            cleanup_sql = f"""
            DELETE FROM {full_target_table}
            WHERE source_uri IN UNNEST(@source_uris)
              AND id NOT IN (
                SELECT id FROM {full_staging_table}
              )
            """
            cleanup_config = bigquery.QueryJobConfig(
                query_parameters=[
                    bigquery.ArrayQueryParameter("source_uris", "STRING", source_uris)
                ]
            )
            cleanup_job = bq_client.query(cleanup_sql, job_config=cleanup_config)
            cleanup_job.result()
            deleted_count = getattr(cleanup_job, "num_dml_affected_rows", 0)
            if isinstance(deleted_count, (int, float)) and deleted_count > 0:
                logger.info("Cleaned up %d orphaned chunks for updated documents.", int(deleted_count))

    finally:
        # 7. Drop Temporary Staging Table
        try:
            bq_client.delete_table(staging_table_id, not_found_ok=True)
            logger.info("Cleaned up temporary staging table %s", staging_table_name)
        except Exception as e:
            logger.warning("Failed to drop staging table %s: %s", staging_table_name, e)

    # 8. Automatically Ensure Vector Index DDL & Monitor Coverage
    ensure_vector_index(bq_client, project_id, dataset_id, table_name)
    check_vector_index_coverage(bq_client, project_id, dataset_id, table_name)

    return len(articles)


def run_test_query(
    query: str,
    project_id: str,
    dataset_id: str,
    table_name: str = "knowledge_articles",
    dry_run: bool = False,
    sample_articles: Optional[list[dict[str, Any]]] = None
):
    """Executes a test vector search query to verify retrieval accuracy."""
    logger.info("--- Testing Query: '%s' ---", query)
    if dry_run or not project_id:
        logger.info("[Dry-Run] Simulating cosine similarity against %d local chunks...", len(sample_articles or []))
        query_vec = generate_text_embedding(query, model_name=DEFAULT_EMBEDDING_MODEL, use_vertex=False)
        scored = []
        for a in (sample_articles or []):
            vec = a.get("embedding") or generate_text_embedding(a["content"], use_vertex=False)
            dot = sum(x * y for x, y in zip(query_vec, vec))
            scored.append((dot, a))
        scored.sort(key=lambda x: x[0], reverse=True)
        for rank, (score, art) in enumerate(scored[:3], 1):
            logger.info("Top %d [%.2f] [%s] %s (ID: %s)", rank, score, art["system"], art["title"], art["id"])
        return

    try:
        from it_helpdesk_agent.tools.enterprise_rag_mcp.knowledge_store import BigQueryVectorKnowledgeStore
        store = BigQueryVectorKnowledgeStore(
            project_id=project_id,
            dataset_id=dataset_id,
            table_name=table_name
        )
        results = store.search(query=query, system="ALL", limit=3)
        for rank, r in enumerate(results, 1):
            logger.info("Top %d [Score: %.2f] [%s] %s (ID: %s)", rank, r.relevance_score, r.system, r.title, r.article_id)
            logger.info("   Snippet: %s", r.snippet[:120])
    except Exception as e:
        logger.error("Test query failed: %s", e)


def main():
    parser = argparse.ArgumentParser(description="Ingest customer documentation into Enterprise Knowledge Base")
    parser.add_argument("--source-dir", type=str, help="Directory containing documents (.md, .txt, .docx, .pdf, .jsonl)")
    parser.add_argument("--file", type=str, help="Single document file to ingest")
    parser.add_argument("--system", type=str, help="Default enterprise system (e.g. ERP, HRM, CRM)")
    parser.add_argument("--project-id", type=str, default=os.getenv("GOOGLE_CLOUD_PROJECT", ""), help="Google Cloud Project ID")
    parser.add_argument("--dataset-id", type=str, default=os.getenv("BIGQUERY_KB_DATASET", "it_helpdesk_kb"), help="BigQuery Dataset ID")
    parser.add_argument("--table-name", type=str, default="knowledge_articles", help="BigQuery Table Name")
    parser.add_argument("--dry-run", action="store_true", help="Parse and embed locally without writing to BigQuery")
    parser.add_argument("--test-query", type=str, help="Run a verification query after ingestion")

    args = parser.parse_args()

    files_to_process: list[Path] = []
    if args.file:
        p = Path(args.file)
        if not p.exists():
            logger.error("File not found: %s", p)
            sys.exit(1)
        files_to_process.append(p)
    elif args.source_dir:
        p = Path(args.source_dir)
        if not p.exists():
            logger.error("Directory not found: %s", p)
            sys.exit(1)
        for ext in ("*.md", "*.txt", "*.docx", "*.pdf", "*.jsonl"):
            files_to_process.extend(p.glob(ext))
    else:
        # Default to data/knowledge_base if present
        default_data_dir = BASE_DIR / "data" / "knowledge_base"
        if default_data_dir.exists():
            for ext in ("*.md", "*.txt", "*.docx", "*.pdf", "*.jsonl"):
                files_to_process.extend(default_data_dir.glob(ext))
        else:
            logger.error("Please specify --source-dir or --file")
            sys.exit(1)

    if not files_to_process:
        logger.warning("No supported document files found to process.")
        sys.exit(0)

    logger.info("Found %d file(s) to process.", len(files_to_process))

    all_articles: list[dict[str, Any]] = []
    for fp in files_to_process:
        logger.info("Parsing %s...", fp.name)
        if fp.suffix.lower() in (".md", ".txt"):
            docs = DocumentParser.parse_markdown_or_text(fp)
        elif fp.suffix.lower() == ".docx":
            docs = DocumentParser.parse_docx(fp)
        elif fp.suffix.lower() == ".pdf":
            docs = DocumentParser.parse_pdf(fp)
        elif fp.suffix.lower() == ".jsonl":
            docs = DocumentParser.parse_jsonl(fp)
        else:
            continue

        for d in docs:
            try:
                processed = process_document(d, default_system=args.system)
                all_articles.extend(processed)
            except Exception as e:
                logger.error("Error processing document from %s: %s", fp.name, e)

    logger.info("Successfully parsed into %d chunked knowledge article(s).", len(all_articles))

    if args.dry_run:
        logger.info("[Dry-Run Mode] Generating sample embeddings locally (No BigQuery writes)...")
        texts = [a["content"] for a in all_articles]
        embeddings = generate_batch_embeddings(texts, model_name=DEFAULT_EMBEDDING_MODEL, use_vertex=False)
        for a, emb in zip(all_articles, embeddings):
            a["embedding"] = emb
        logger.info("[Dry-Run Mode] All %d articles validated and embedded successfully.", len(all_articles))
    else:
        if not args.project_id:
            logger.error("Project ID is required for BigQuery ingestion. Set GOOGLE_CLOUD_PROJECT or pass --project-id")
            sys.exit(1)
        ingest_articles_to_bigquery(
            all_articles,
            project_id=args.project_id,
            dataset_id=args.dataset_id,
            table_name=args.table_name
        )

    if args.test_query:
        run_test_query(
            args.test_query,
            project_id=args.project_id,
            dataset_id=args.dataset_id,
            table_name=args.table_name,
            dry_run=args.dry_run,
            sample_articles=all_articles
        )


if __name__ == "__main__":
    main()
