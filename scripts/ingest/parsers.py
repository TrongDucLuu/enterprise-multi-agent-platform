"""
Document Parsers for Enterprise Knowledge Base.
Extracts raw text, structured sections, and hierarchy from Markdown, Text, DOCX, PDF, and JSONL.
"""

import re
import time
import json
import logging
from pathlib import Path
from typing import Optional, Any

from it_helpdesk_agent.app_utils.system_config import get_document_processing_config

PARSER_VERSION = "1.0.0"

logger = logging.getLogger("ingest.parsers")


class DocumentParser:
    """Extracts raw text, structured sections, and metadata from various enterprise document formats."""

    @staticmethod
    def parse_markdown_or_text(file_path: Path) -> list[dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        title = file_path.stem.replace("_", " ").title()
        owner: Optional[str] = None
        effective_date: Optional[str] = None
        expiry_date: Optional[str] = None

        # Check for optional YAML frontmatter
        if content.startswith("---"):
            fm_match = re.match(r"^---\s*\n(.*?)\n---\s*\n(.*)$", content, re.DOTALL)
            if fm_match:
                fm_text = fm_match.group(1)
                content = fm_match.group(2)
                for line in fm_text.splitlines():
                    if ":" in line:
                        k, v = line.split(":", 1)
                        k_clean = k.strip().lower()
                        v_clean = v.strip().strip("'\"")
                        if k_clean == "title" and v_clean:
                            title = v_clean
                        elif k_clean == "owner" and v_clean:
                            owner = v_clean
                        elif k_clean == "effective_date" and v_clean:
                            effective_date = v_clean
                        elif k_clean == "expiry_date" and v_clean:
                            expiry_date = v_clean

        heading_pattern = re.compile(r"^(#{1,3})\s+(.+)$")

        sections: list[dict[str, Any]] = []
        current_heading = title
        current_level = 1
        current_h1: Optional[str] = title
        current_h2: Optional[str] = None
        current_h3: Optional[str] = None
        current_lines: list[str] = []

        for line in content.splitlines():
            m = heading_pattern.match(line)
            if m:
                if current_lines:
                    sec_text = "\n".join(current_lines).strip()
                    if sec_text:
                        sections.append({
                            "level": current_level,
                            "heading": current_heading,
                            "content": sec_text,
                            "hierarchy": {
                                "h1": current_h1,
                                "h2": current_h2,
                                "h3": current_h3,
                            },
                        })
                    current_lines = []
                current_level = len(m.group(1))
                current_heading = m.group(2).strip()
                if current_level == 1:
                    current_h1 = current_heading
                    current_h2 = None
                    current_h3 = None
                    if not sections:
                        title = current_heading
                elif current_level == 2:
                    current_h2 = current_heading
                    current_h3 = None
                elif current_level == 3:
                    current_h3 = current_heading
            else:
                current_lines.append(line)

        if current_lines:
            sec_text = "\n".join(current_lines).strip()
            if sec_text:
                sections.append({
                    "level": current_level,
                    "heading": current_heading,
                    "content": sec_text,
                    "hierarchy": {
                        "h1": current_h1,
                        "h2": current_h2,
                        "h3": current_h3,
                    },
                })

        return [{
            "title": title,
            "content": content,
            "source_uri": str(file_path),
            "file_type": file_path.suffix.lower(),
            "sections": sections,
            "owner": owner,
            "effective_date": effective_date,
            "expiry_date": expiry_date,
        }]

    @staticmethod
    def parse_docx(file_path: Path) -> list[dict[str, Any]]:
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            title = paragraphs[0] if paragraphs else file_path.stem.replace("_", " ").title()

            sections: list[dict[str, Any]] = []
            current_heading = title
            current_level = 1
            current_h1: Optional[str] = title
            current_h2: Optional[str] = None
            current_h3: Optional[str] = None
            current_paras: list[str] = []

            for p in doc.paragraphs:
                p_text = p.text.strip()
                if not p_text:
                    continue
                style_name = getattr(p.style, "name", "") or ""
                if "heading" in style_name.lower():
                    if current_paras:
                        sec_text = "\n\n".join(current_paras).strip()
                        if sec_text:
                            sections.append({
                                "level": current_level,
                                "heading": current_heading,
                                "content": sec_text,
                                "hierarchy": {
                                    "h1": current_h1,
                                    "h2": current_h2,
                                    "h3": current_h3,
                                },
                            })
                        current_paras = []
                    m = re.search(r"\d+", style_name)
                    current_level = int(m.group(0)) if m else 1
                    current_heading = p_text
                    if current_level == 1:
                        current_h1 = current_heading
                        current_h2 = None
                        current_h3 = None
                        if not sections:
                            title = current_heading
                    elif current_level == 2:
                        current_h2 = current_heading
                        current_h3 = None
                    elif current_level == 3:
                        current_h3 = current_heading
                else:
                    current_paras.append(p_text)

            if current_paras:
                sec_text = "\n\n".join(current_paras).strip()
                if sec_text:
                    sections.append({
                        "level": current_level,
                        "heading": current_heading,
                        "content": sec_text,
                        "hierarchy": {
                            "h1": current_h1,
                            "h2": current_h2,
                            "h3": current_h3,
                        },
                    })

            return [{
                "title": title,
                "content": content,
                "source_uri": str(file_path),
                "file_type": ".docx",
                "sections": sections,
            }]
        except ImportError:
            logger.warning("python-docx is not installed. To parse DOCX files, install python-docx (`pip install python-docx`).")
            return []
        except Exception as e:
            logger.error("Failed to parse docx file %s: %s", file_path, e)
            return []

    @staticmethod
    def parse_pdf_document_ai(file_path: Path, config: dict[str, Any]) -> list[dict[str, Any]]:
        """
        Parses a PDF using Google Cloud Document AI Layout Parser.
        Extracts structured sections (heading-1, heading-2, paragraphs, tables).
        Retries up to max_retries with exponential backoff and timeout.
        Fails closed on unrecoverable errors.
        """
        processor_id = config.get("document_ai_processor_id")
        if not processor_id:
            raise RuntimeError("Missing 'document_ai_processor_id' for Document AI PDF parsing. (Fail-Closed)")

        timeout_seconds = config.get("document_ai_timeout_seconds", 60.0)
        max_retries = config.get("document_ai_max_retries", 2)

        try:
            from google.cloud import documentai
        except ImportError:
            logger.error("google-cloud-documentai is not installed. Run `pip install google-cloud-documentai`.")
            raise

        with open(file_path, "rb") as f:
            pdf_bytes = f.read()

        raw_document = documentai.RawDocument(content=pdf_bytes, mime_type="application/pdf")
        request = documentai.ProcessRequest(name=processor_id, raw_document=raw_document)

        client = documentai.DocumentProcessorServiceClient()

        last_error = None
        for attempt in range(max_retries + 1):
            try:
                logger.info(
                    "Calling Document AI Layout Parser for %s (attempt %d/%d, timeout=%.1fs)...",
                    file_path.name, attempt + 1, max_retries + 1, timeout_seconds
                )
                result = client.process_document(request=request, timeout=timeout_seconds)
                doc = result.document
                return DocumentParser._map_document_ai_to_document_info(doc, file_path)
            except Exception as e:
                last_error = e
                logger.warning("Document AI attempt %d failed for %s: %s", attempt + 1, file_path.name, e)
                if attempt < max_retries:
                    backoff = 2 ** attempt
                    time.sleep(backoff)

        raise RuntimeError(
            f"Document AI parsing failed for '{file_path}' after {max_retries + 1} attempts: {last_error}. (Fail-Closed)"
        )

    @staticmethod
    def _map_document_ai_to_document_info(doc: Any, file_path: Path) -> list[dict[str, Any]]:
        full_text = getattr(doc, "text", None) if not isinstance(doc, dict) else doc.get("text", "")
        if full_text is None:
            full_text = ""

        blocks = []
        pages = getattr(doc, "pages", []) if not isinstance(doc, dict) else doc.get("pages", [])
        for page in pages:
            page_blocks = getattr(page, "blocks", []) if not isinstance(page, dict) else page.get("blocks", [])
            for b in page_blocks:
                blocks.append(b)

        title = file_path.stem.replace("_", " ").title()
        sections: list[dict[str, Any]] = []
        current_heading = title
        current_level = 1
        current_h1: Optional[str] = title
        current_h2: Optional[str] = None
        current_h3: Optional[str] = None
        current_content_parts: list[str] = []

        def get_block_type_and_text(block: Any) -> tuple[str, str]:
            if isinstance(block, dict):
                b_type = block.get("type_") or block.get("block_type") or block.get("type", "paragraph")
                b_text = block.get("text", "")
                if not b_text and "layout" in block:
                    anchors = block["layout"].get("text_anchor", {}).get("text_segments", [])
                    b_text = "".join(full_text[int(seg.get("start_index", 0)):int(seg.get("end_index", 0))] for seg in anchors)
                return str(b_type), b_text.strip()
            else:
                b_type = getattr(block, "type_", None) or getattr(block, "block_type", "paragraph")
                b_text = getattr(block, "text", "")
                if not b_text and hasattr(block, "layout") and hasattr(block.layout, "text_anchor"):
                    segments = getattr(block.layout.text_anchor, "text_segments", [])
                    b_text = "".join(full_text[int(getattr(s, "start_index", 0)):int(getattr(s, "end_index", 0))] for s in segments)
                return str(b_type), b_text.strip()

        for b in blocks:
            b_type, b_text = get_block_type_and_text(b)
            if not b_text:
                continue

            if "heading" in b_type.lower():
                if current_content_parts:
                    sec_text = "\n\n".join(current_content_parts).strip()
                    if sec_text:
                        sections.append({
                            "level": current_level,
                            "heading": current_heading,
                            "content": sec_text,
                            "hierarchy": {
                                "h1": current_h1,
                                "h2": current_h2,
                                "h3": current_h3,
                            },
                        })
                    current_content_parts = []

                m = re.search(r"\d+", b_type)
                current_level = int(m.group(0)) if m else 1
                current_heading = b_text
                if current_level == 1:
                    current_h1 = current_heading
                    current_h2 = None
                    current_h3 = None
                    if not sections:
                        title = current_heading
                elif current_level == 2:
                    current_h2 = current_heading
                    current_h3 = None
                elif current_level == 3:
                    current_h3 = current_heading
            else:
                current_content_parts.append(b_text)

        if current_content_parts:
            sec_text = "\n\n".join(current_content_parts).strip()
            if sec_text:
                sections.append({
                    "level": current_level,
                    "heading": current_heading,
                    "content": sec_text,
                    "hierarchy": {
                        "h1": current_h1,
                        "h2": current_h2,
                        "h3": current_h3,
                    },
                })

        return [{
            "title": title,
            "content": full_text.strip() if full_text.strip() else "\n\n".join(s["content"] for s in sections),
            "source_uri": str(file_path),
            "file_type": ".pdf",
            "sections": sections,
        }]

    @staticmethod
    def parse_pdf(file_path: Path, doc_proc_config: Optional[dict[str, Any]] = None) -> list[dict[str, Any]]:
        if doc_proc_config is None:
            try:
                doc_proc_config = get_document_processing_config()
            except Exception as e:
                logger.debug("Could not load document processing config: %s", e)
                doc_proc_config = {"pdf_parser": "pypdf_flat"}

        pdf_parser_mode = (doc_proc_config or {}).get("pdf_parser", "pypdf_flat")
        if pdf_parser_mode == "document_ai":
            return DocumentParser.parse_pdf_document_ai(file_path, doc_proc_config)

        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text.strip())
            content = "\n\n".join(pages_text)
            title = file_path.stem.replace("_", " ").title()
            if pages_text:
                first_lines = [l.strip() for l in pages_text[0].splitlines() if l.strip()]
                if first_lines:
                    title = first_lines[0][:100]
            return [{
                "title": title,
                "content": content,
                "source_uri": str(file_path),
                "file_type": ".pdf",
                "sections": [],
            }]
        except ImportError:
            logger.warning("pypdf is not installed. To parse PDF files, install pypdf (`pip install pypdf`).")
            return []
        except Exception as e:
            logger.error("Failed to parse pdf file %s: %s", file_path, e)
            return []

    @staticmethod
    def parse_jsonl(file_path: Path) -> list[dict[str, Any]]:
        articles = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line_no, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    source_uri = data.get("source_uri") or f"{file_path}#L{line_no}"
                    articles.append({
                        "id": data.get("id"),
                        "system": data.get("system"),
                        "title": data.get("title", f"Article {line_no}"),
                        "category": data.get("category", "General"),
                        "content": data.get("content", ""),
                        "keywords": data.get("keywords", []),
                        "source_uri": source_uri,
                        "file_type": ".jsonl",
                        "sections": [],
                        "owner": data.get("owner"),
                        "effective_date": data.get("effective_date"),
                        "expiry_date": data.get("expiry_date"),
                    })
                except json.JSONDecodeError as e:
                    logger.warning("Invalid JSON at line %d in %s: %s", line_no, file_path, e)
        return articles
